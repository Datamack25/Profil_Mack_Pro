import streamlit as st
import plotly.graph_objects as go
import pandas as pd
import requests
import io
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY

# ─── PAGE CONFIG ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Mackenson Cineus — Finance Professional",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CUSTOM CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
  @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700;900&family=DM+Sans:wght@300;400;500;600&display=swap');
  :root {
    --gold: #C9A84C; --gold-light: #E8C97A;
    --navy: #0D1B2A; --navy-mid: #1A2E42; --navy-light: #243B55;
    --cream: #F5F0E8; --white: #FFFFFF;
    --accent: #4A9EFF; --success: #2ECC71; --warn: #E67E22;
  }
  .stApp { background: var(--navy); color: var(--cream); }
  .stApp > header { background: transparent !important; }
  [data-testid="stSidebar"] {
    background: linear-gradient(180deg, var(--navy-mid) 0%, var(--navy) 100%);
    border-right: 1px solid rgba(201,168,76,0.2);
  }
  [data-testid="stSidebar"] * { color: var(--cream) !important; }
  h1,h2,h3 { font-family: 'Playfair Display', serif !important; }
  p,div,span,label { font-family: 'DM Sans', sans-serif !important; }

  .hero-banner {
    background: linear-gradient(135deg, var(--navy-mid) 0%, var(--navy-light) 50%, var(--navy-mid) 100%);
    border: 1px solid rgba(201,168,76,0.3); border-radius: 16px;
    padding: 40px; margin-bottom: 32px; position: relative; overflow: hidden;
  }
  .hero-banner::before {
    content:''; position:absolute; top:-50%; right:-10%; width:400px; height:400px;
    background: radial-gradient(circle, rgba(201,168,76,0.08) 0%, transparent 70%); pointer-events:none;
  }
  .hero-name {
    font-family: 'Playfair Display', serif !important; font-size: 3rem !important; font-weight: 900 !important;
    background: linear-gradient(135deg, var(--gold) 0%, var(--gold-light) 100%);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; line-height: 1.1; margin: 0;
  }
  .hero-title { font-size:1.05rem; color:var(--cream); opacity:0.85; margin-top:8px; letter-spacing:2px; text-transform:uppercase; }
  .hero-badge {
    display:inline-block; background:rgba(201,168,76,0.15); border:1px solid rgba(201,168,76,0.4);
    border-radius:50px; padding:4px 14px; font-size:0.8rem; color:var(--gold-light); margin:4px;
  }
  .metric-card {
    background:linear-gradient(135deg,var(--navy-mid),var(--navy-light));
    border:1px solid rgba(201,168,76,0.2); border-radius:12px; padding:24px; text-align:center;
  }
  .metric-value { font-family:'Playfair Display',serif !important; font-size:2.5rem; font-weight:700; color:var(--gold); display:block; }
  .metric-label { font-size:0.85rem; color:var(--cream); opacity:0.7; text-transform:uppercase; letter-spacing:1px; }
  .section-header {
    font-family:'Playfair Display',serif !important; font-size:1.8rem !important; color:var(--gold) !important;
    border-bottom:1px solid rgba(201,168,76,0.3); padding-bottom:12px; margin-bottom:24px;
  }

  /* Bio periods */
  .bio-period {
    background:var(--navy-mid); border:1px solid rgba(201,168,76,0.15);
    border-radius:16px; padding:28px 32px; margin-bottom:24px;
    position:relative; overflow:hidden;
  }
  .bio-period::before { content:''; position:absolute; left:0; top:0; bottom:0; width:4px; }
  .bio-period.phase1::before { background:linear-gradient(180deg,#4A9EFF,#2ECC71); }
  .bio-period.phase2::before { background:linear-gradient(180deg,#C9A84C,#E8C97A); }
  .bio-period.phase3::before { background:linear-gradient(180deg,#E67E22,#C9A84C); }
  .bio-period.phase4::before { background:linear-gradient(180deg,#9B59B6,#4A9EFF); }
  .bio-phase-label { font-size:0.72rem; letter-spacing:2px; text-transform:uppercase; font-weight:600; margin-bottom:6px; }
  .phase1 .bio-phase-label { color:#4A9EFF; }
  .phase2 .bio-phase-label { color:var(--gold); }
  .phase3 .bio-phase-label { color:var(--warn); }
  .phase4 .bio-phase-label { color:#9B59B6; }
  .bio-period-title { font-family:'Playfair Display',serif !important; font-size:1.3rem; font-weight:700; color:var(--white); margin-bottom:8px; }
  .bio-period-years { font-size:0.8rem; color:var(--cream); opacity:0.55; margin-bottom:14px; letter-spacing:1px; }
  .bio-period-text { font-size:0.95rem; line-height:1.85; color:var(--cream); opacity:0.88; }
  .bio-tag {
    display:inline-block; background:rgba(255,255,255,0.06); border:1px solid rgba(255,255,255,0.12);
    border-radius:4px; padding:2px 10px; font-size:0.75rem; color:var(--cream); opacity:0.7; margin:3px 2px 0;
  }

  .timeline-item {
    background:var(--navy-mid); border-left:3px solid var(--gold);
    border-radius:0 12px 12px 0; padding:16px 20px; margin-bottom:14px;
  }
  .timeline-item:hover { background:var(--navy-light); }
  .timeline-year { color:var(--gold); font-weight:600; font-size:0.85rem; letter-spacing:1px; }
  .timeline-title { font-size:1.05rem; font-weight:600; color:var(--white); margin:4px 0; }
  .timeline-desc { font-size:0.9rem; color:var(--cream); opacity:0.8; }

  .skill-bar-container { margin-bottom:12px; }
  .skill-name { font-size:0.9rem; color:var(--cream); margin-bottom:4px; }
  .skill-bar { height:8px; background:rgba(255,255,255,0.1); border-radius:4px; overflow:hidden; }
  .skill-fill { height:100%; border-radius:4px; background:linear-gradient(90deg,var(--gold),var(--gold-light)); }

  .platform-card {
    background:var(--navy-mid); border:1px solid rgba(255,255,255,0.08);
    border-radius:12px; padding:16px 20px; margin-bottom:12px;
    display:flex; align-items:flex-start; gap:14px; transition:border-color 0.3s,background 0.3s;
  }
  .platform-card:hover { border-color:rgba(201,168,76,0.4); background:var(--navy-light); }
  .platform-icon { font-size:1.6rem; flex-shrink:0; width:44px; height:44px; display:flex; align-items:center; justify-content:center; background:rgba(201,168,76,0.1); border-radius:10px; }
  .platform-name { font-size:0.95rem; font-weight:600; color:var(--gold); margin-bottom:2px; }
  .platform-desc { font-size:0.82rem; color:var(--cream); opacity:0.7; line-height:1.5; }
  .platform-link { display:inline-block; margin-top:6px; font-size:0.78rem; color:var(--accent); text-decoration:none; border-bottom:1px solid rgba(74,158,255,0.3); }

  .stButton > button {
    background:linear-gradient(135deg,var(--gold) 0%,var(--gold-light) 100%) !important;
    color:var(--navy) !important; border:none !important; border-radius:8px !important;
    font-weight:600 !important; letter-spacing:0.5px; padding:10px 24px !important;
  }
  .stButton > button:hover { box-shadow:0 8px 24px rgba(201,168,76,0.3) !important; }
  .stTabs [data-baseweb="tab-list"] { background:var(--navy-mid); border-radius:8px; gap:4px; }
  .stTabs [data-baseweb="tab"] { color:var(--cream) !important; }
  .stTabs [aria-selected="true"] { background:rgba(201,168,76,0.2) !important; color:var(--gold) !important; }
  .stTextArea textarea,.stTextInput input,.stSelectbox select {
    background:var(--navy-mid) !important; color:var(--cream) !important;
    border:1px solid rgba(201,168,76,0.3) !important; border-radius:8px !important;
  }
  .generated-doc {
    background:var(--navy-mid); border:1px solid rgba(201,168,76,0.2);
    border-radius:12px; padding:32px; line-height:1.8; white-space:pre-wrap; color:var(--cream);
  }
  .doc-header {
    font-family:'Playfair Display',serif; font-size:1.4rem; color:var(--gold);
    margin-bottom:16px; border-bottom:1px solid rgba(201,168,76,0.2); padding-bottom:12px;
  }
  #MainMenu,footer,[data-testid="stHeader"] { display:none !important; }
  .block-container { padding-top:2rem !important; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# DATA
# ═══════════════════════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════════════════════
# DATA — SOURCE UNIQUE DE VÉRITÉ (modifiable depuis la page "Édition")
# ═══════════════════════════════════════════════════════════════════════════════

# Initialisation session_state pour l'édition en ligne


def _init_state():
    defaults = {
        # ── IDENTITÉ ──────────────────────────────────────────────────────────
        "edit_name":     "Mackenson CINÉUS",
        "edit_title":    "Analyste LCB-FT | Compliance Officer",
        "edit_location": "Île-de-France, France",
        "edit_email":    "mackenson.cineus@email.com",
        "edit_phone":    "+33 6 XX XX XX XX",
        "edit_linkedin": "linkedin.com/in/mackenson-cineus",
        "edit_summary_fr": (
            "Analyste LCB-FT en alternance chez Banque Delubac & Cie sur clientèle corporate, CIB, "
            "asset management et banque privée. Expérience Chargé de Conformité chez HiPay SAS "
            "(fintech agréée ACPR) : risques pays GAFI, KYC/PPE/SOE, pays sensibles. "
            "MBA 2 Finance de Marché en cours (ESLSCA Paris, parcours anglais). "
            "Certifié AMF, CAMS (Sanctions & Embargos), candidat CFA Level I. "
            "Lauréat Hackathon Fintech Générations 2023 — France FinTech / Société Générale."
        ),
        "edit_summary_en": (
            "Work-study Analyst at Banque Delubac & Cie with expertise in transaction monitoring "
            "for Corporate, CIB, and Asset Management clients. Proven experience in Fintech "
            "Compliance (HiPay SAS) focusing on FATF risk countries, KYC/PEP/SOE, and "
            "high-risk jurisdictions. Currently pursuing an MBA in Market Finance at ESLSCA Paris. "
            "AMF Certified, CAMS (Sanctions & Embargoes), and CFA Level I Candidate."
        ),
        "edit_memoir_title": (
            "Impacts of ESG Factors and Regulatory Constraints on Portfolio Management"
        ),
        "edit_memoir_desc": (
            "80 pages · Rédigé en anglais · MBA 2 Finance de marché · ESLSCA Paris. "
            "Performance des portefeuilles sous contrainte réglementaire ESG/non-ESG — "
            "analyse du risque et du rendement."
        ),

        # ── EXPÉRIENCES ───────────────────────────────────────────────────────
        "edit_exp": [
            {
                "role":     "Analyste LCB-FT",
                "role_en":  "AML/CFT Analyst",
                "org":      "Banque DELUBAC & CIE · Paris, France",
                "org_type": "Banque",
                "period":   "Septembre 2025 – Aujourd'hui",
                "period_en":"September 2025 – Present",
                "url":      "https://www.delubac.com",
                "bullets_fr": [
                    "Traitement des alertes de niveau 2 LCB-FT sur la clientèle (corporate, CIB, asset management, banque privée) — analyse de la situation financière, économique et juridique",
                    "Déclaration de soupçons : analyse, validation et création de règles relatives aux demandes d'exclusions sur certaines transactions récurrentes",
                    "Analyse KYC, identification des bénéficiaires effectifs et traitement des Examens Renforcés (ER)",
                    "Analyser les états financiers des clients pour évaluer leur santé financière et la cohérence des transactions",
                    "Traiter et escalader les demandes de droit de communication et les réquisitions judiciaires",
                ],
                "bullets_en": [
                    "Level 2 AML/CFT alert processing for corporate, CIB, asset management and private banking clients — financial, economic and legal situation analysis",
                    "Suspicious activity reports (SAR): analysis, validation and creation of exclusion rules for recurring transactions",
                    "KYC analysis, beneficial owner identification and Enhanced Due Diligence (EDD) processing",
                    "Financial statement analysis to assess client financial health and transaction consistency",
                    "Processing and escalating judicial requisitions and right-of-communication requests",
                ],
            },
            {
                "role":     "Chargé de conformité",
                "role_en":  "Compliance Officer",
                "org":      "HIPAY SAS · Nantes-Paris, France",
                "org_type": "Fintech",
                "period":   "Septembre 2024 – Septembre 2025",
                "period_en":"September 2024 – September 2025",
                "url":      "https://hipay.com",
                "bullets_fr": [
                    "Analyser les pays à risque (GAFI) et les typologies de fraude et de blanchiment (LCB-FT) sur pays sensibles et complexes ; traitement des alertes sur la clientèle du secteur des jeux en ligne et des places de marchés",
                    "Identifier les bénéficiaires effectifs, les personnes politiquement exposées (PPE) et les entreprises détenues par l'État",
                    "Traiter les alertes et répondre aux réquisitions judiciaires ; réaliser des analyses complémentaires (fraudes, sanctions, demandes judiciaires)",
                    "Due Diligence des partenaires institutionnels et établissements financiers systémiques",
                    "Mise à jour de la documentation de risques (cartographie des risques) et participation aux audits internes",
                ],
                "bullets_en": [
                    "FATF risk country analysis and AML/CFT fraud typologies on sensitive and complex jurisdictions; alert processing for online gaming and marketplace clients",
                    "Beneficial owner identification, Politically Exposed Persons (PEP) screening and State-Owned Entities (SOE) identification",
                    "Alert processing and judicial requisition responses; complementary fraud and sanctions analysis",
                    "Institutional partner Due Diligence and systemic financial institution screening",
                    "Risk documentation update (risk mapping) and participation in internal compliance audits",
                ],
            },
            {
                "role":     "Lauréat 1er Prix — Hackathon Fintech Générations 2023",
                "role_en":  "1st Prize Winner — Fintech Générations Hackathon 2023",
                "org":      "France FinTech · Société Générale · Treezor · Paris&Co",
                "org_type": "Compétition",
                "period":   "Octobre 2023",
                "period_en":"October 2023",
                "url":      "https://francefintech.org",
                "bullets_fr": [
                    "Projet Victoria : solution fintech de financement des rénovations DPE par financement privé — conception et développement en 48h devant jury d'experts",
                    "Pitch à Fintech R:Evolution (1 500 participants : investisseurs, entrepreneurs, régulateurs) — invitation et présentation devant fonds d'investissement",
                ],
                "bullets_en": [
                    "Project Victoria: fintech solution for DPE renovation financing through private funding — designed and developed in 48h before expert jury",
                    "Pitch at Fintech R:Evolution (1,500 participants: investors, entrepreneurs, regulators) — invitation and presentation before investment funds",
                ],
            },
            {
                "role":     "Board Member",
                "role_en":  "Board Member",
                "org":      "Erasmus Expertise · International",
                "org_type": "ONG",
                "period":   "2021 – Présent",
                "period_en":"2021 – Present",
                "url":      "https://erasmus-expertise.eu",
                "bullets_fr": [
                    "Gouvernance ONG internationale : développement durable, éducation, inclusion sociale",
                    "Accompagnement linguistique et administratif de 7 étudiants internationaux — mise en place groupe de conversation français",
                ],
                "bullets_en": [
                    "International NGO governance: sustainable development, education, social inclusion",
                    "Linguistic and administrative support for 7 international students — setup of French conversation group",
                ],
            },
            {
                "role":     "Directeur Financier & Directeur des Opérations",
                "role_en":  "CFO & Director of Operations",
                "org":      "Hult Prize Haïti · Haïti / République Dominicaine",
                "org_type": "Association",
                "period":   "2018 – 2021",
                "period_en":"2018 – 2021",
                "url":      "https://www.hultprize.org",
                "bullets_fr": [
                    "Directeur financier : recherche de financements (Ministère des Finances, BNC), élaboration du budget de l'association, supervision des responsables régionaux dans 10 départements",
                    "Représentant Haïti au Hult Prize international 2018 — projet WEISS (biogaz & fertilisant organique) ; volontariat international en République Dominicaine",
                ],
                "bullets_en": [
                    "CFO: fundraising (Ministry of Finance, BNC), budget preparation, supervision of regional managers across 10 departments",
                    "Represented Haiti at Hult Prize International 2018 — WEISS project (biogas & organic fertilizer); international volunteering in Dominican Republic",
                ],
            },
        ],

        # ── FORMATION ────────────────────────────────────────────────────────
        "edit_edu": [
            {
                "deg":     "MBA 2 Finance de Marché (Parcours en anglais)",
                "deg_en":  "MBA in Market Finance (English Track)",
                "school":  "ESLSCA Business School Paris · Paris, France",
                "yr":      "Septembre 2024 – Septembre 2026",
                "yr_en":   "September 2024 – September 2026",
                "url":     "https://www.eslsca.fr",
                "det":     "Audit financier · Comptabilité approfondie · Gestion des risques · Gestion de Portefeuille · Pricing d'option · Produits dérivés · Investissements alternatifs · Analyse financière · Conformité · Évaluation d'actifs · Machine learning · VBA & Python",
                "det_en":  "Financial Audit · Advanced Accounting · Risk Management · Portfolio Management · Options Pricing · Derivatives · Alternative Investments · Financial Analysis · Compliance · Asset Valuation · Machine Learning · VBA & Python",
                "mention": "",
            },
            {
                "deg":     "Licence en Banque, Finance et Assurance",
                "deg_en":  "Bachelor in Banking, Finance and Insurance",
                "school":  "Université du Mans · Le Mans, France",
                "yr":      "Septembre 2022 – Juin 2024",
                "yr_en":   "September 2022 – June 2024",
                "url":     "https://www.univ-lemans.fr",
                "det":     "Analyse et gestion de portefeuille · Droit commercial · Droit fiscal · VBA & analyse stochastique · Comptabilité de gestion · Mathématiques financières · Économétrie · Assurance vie & non-vie · Gestion des risques bancaires",
                "det_en":  "Portfolio Analysis & Management · Commercial Law · Tax Law · VBA & Stochastic Analysis · Management Accounting · Financial Mathematics · Econometrics · Life & Non-Life Insurance · Banking Risk Management",
                "mention": "",
            },
        ],

        # ── COMPÉTENCES ───────────────────────────────────────────────────────
        "edit_skills": {
            "LCB-FT / AML / AMLD5-6":          95,
            "KYC / KYB / Due Diligence":         92,
            "Analyse risques pays (GAFI)":        90,
            "Déclarations de soupçon (DS/SAR)":  88,
            "Réglementation ACPR / GAFI":         90,
            "Marchés Financiers & Trading":       82,
            "Analyse financière (états financiers)": 80,
            "Excel avancé / VBA / Python":        75,
            "Bloomberg / Looker / Jura":          68,
        },
        "edit_certifs": [
            "AMF",
            "CAMS – Sanctions & Embargos",
            "CFA Level I (en cours de préparation)",
        ],
        "edit_tech": "Excel avancé – VBA – Python – Pack Microsoft Office – Looker – Jura",
        "edit_interests": "Sécurité financière – Création de modèles financiers – Programmation",
        "edit_distinctions": [
            ("🥇", "Hackathon Fintech Générations 2023",
             "1er Prix — Projet Victoria (DPE) · France FinTech / Société Générale"),
            ("🏆", "Hult Prize 2018",
             "Représentant Haïti — Projet WEISS (biogaz & fertilisant organique)"),
            ("🏅", "Hackathon Unleash (ODD)",
             "Lauréat — Objectifs de Développement Durable — équipe internationale"),
            ("🎙", "Podcast INCLUTECH",
             "Créateur & animateur — Finance & Inclusion — Spotify & Apple Podcasts"),
        ],
        "edit_personality": (
            "Analyste LCB-FT rigoureux à l'esprit entrepreneurial — "
            "de Port-au-Prince aux salles de conformité parisiennes, "
            "il transforme chaque contrainte réglementaire en levier de valeur."
        ),
        "edit_phone":    "+33 6 XX XX XX XX",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init_state()

def _P(lang="fr"):
    ss = st.session_state
    return {
        "name":     ss.get("edit_name",    "Mackenson CINÉUS"),
        "title":    ss.get("edit_title",   "Analyste LCB-FT | Compliance Officer"),
        "location": ss.get("edit_location","Île-de-France, France"),
        "email":    ss.get("edit_email",   "mackenson.cineus@email.com"),
        "phone":    ss.get("edit_phone",   "+33 6 XX XX XX XX"),
        "linkedin": ss.get("edit_linkedin","linkedin.com/in/mackenson-cineus"),
        "summary":  ss.get(f"edit_summary_{lang}", ss.get("edit_summary_fr","")),
    }

PROFILE = _P()

# ── TIMELINE ──────────────────────────────────────────────────────────────────
TIMELINE = [
    {"year":"2016–2018","title":"Président — Groupe Économie & Finance · ACTIVEH","desc":"Leadership 15+ bénévoles · Projet Goud la se Pa m · conférences économie & entrepreneuriat","url":"","icon":"🤝","cat":"Social"},
    {"year":"2018–2021","title":"Directeur Financier & Opérations — Hult Prize Haïti","desc":"Financement · budget · supervision 10 dép. · volontariat RD · Hult Prize 2018 (WEISS)","url":"https://www.hultprize.org","icon":"🏆","cat":"Entrepreneuriat"},
    {"year":"2018–2021","title":"Directeur Exécutif & Fondateur — ANGAJMAN","desc":"Projets citoyens 10 dép. Haïti · campagne COVID-19 · 50 jeunes · initiatives ESG","url":"","icon":"🌱","cat":"Social"},
    {"year":"2018–2021","title":"PHJEA — Section Finance","desc":"Parlement Haïtien Jeunesse Eau & Assainissement · Parlement Mondial de l'Eau","url":"","icon":"🌊","cat":"Social"},
    {"year":"2021","title":"Hackathon Unleash (ODD) — Lauréat","desc":"Objectifs de Développement Durable · équipe internationale · rôle mentor","url":"","icon":"🌍","cat":"Entrepreneuriat"},
    {"year":"2022–2024","title":"Licence BFA — Université du Mans","desc":"Banque, Finance & Assurance · VBA · analyse stochastique · droit fiscal","url":"https://www.univ-lemans.fr","icon":"🏫","cat":"Éducation"},
    {"year":"2021–Présent","title":"Board Member — Erasmus Expertise","desc":"ONG internationale · tuteur 7 étudiants · groupe conversation · CA 2023","url":"https://erasmus-expertise.eu","icon":"🌍","cat":"Réseau"},
    {"year":"2021–Présent","title":"Podcast INCLUTECH","desc":"Finance & Inclusion Financière · bi-mensuel · Spotify & Apple Podcasts","url":"https://open.spotify.com/show/5XvFdWYwhHWY3EguIhhf69","icon":"🎙","cat":"Réseau"},
    {"year":"2024–Présent","title":"MBA 2 Finance de Marché — ESLSCA Paris","desc":"Parcours anglais · Produits dérivés · Conformité · VBA & Python · Machine learning","url":"https://www.eslsca.fr","icon":"📈","cat":"Éducation"},
    {"year":"Oct. 2023","title":"🥇 Hackathon Fintech Générations — 1er Prix","desc":"Projet Victoria (DPE) · France FinTech / SG / Treezor · Fintech R:Evolution 1 500 participants","url":"https://francefintech.org","icon":"🚀","cat":"Entrepreneuriat"},
    {"year":"2024–2025","title":"Chargé de conformité — HIPAY SAS","desc":"LCB-FT · Pays sensibles · PPE/SOE · Tracfin · Jeux en ligne · Marketplaces","url":"https://hipay.com","icon":"💳","cat":"Professionnel"},
    {"year":"2025–Présent","title":"Analyste LCB-FT — Banque DELUBAC & CIE","desc":"Alertes N2 · ER · KYC corporate/CIB/PSP · DS/SAR · Réquisitions judiciaires","url":"https://www.delubac.com","icon":"🏦","cat":"Professionnel"},
]

SKILLS = dict(st.session_state.get("edit_skills", {
    "LCB-FT / AML": 95, "KYC/KYB": 92, "Risques pays": 90,
    "DS/SAR": 88, "ACPR/GAFI": 90, "Marchés Financiers": 82,
}))

LANGUAGES = {"Français (Natif)": 100, "Anglais (C1)": 88, "Créole haïtien": 100}

GEO_DATA = {
    "locations": ["Port-au-Prince, Haïti", "Le Mans, France", "Paris / Levallois-Perret"],
    "lat": [18.5432, 47.9960, 48.8924],
    "lon": [-72.3388, 0.1966, 2.2873],
    "events": ["ACTIVEH · Hult Prize · ANGAJMAN · PHJEA · Unleash",
               "Licence BFA · Université du Mans",
               "MBA ESLSCA · Delubac LCB-FT · HiPay · France FinTech"],
    "years": ["2014–2021", "2022–2024", "2024–Présent"],
    "size": [18, 14, 34],
}

BIO_PERIODS = [
    {
        "phase": "phase1",
        "label": "Phase 1 · 2014–2021 · Haïti",
        "title": "Racines académiques, leadership associatif & entrepreneuriat social",
        "years": "INAGHEI · ACTIVEH · Hult Prize · ANGAJMAN · PHJEA · Unleash",
        "text": [
            'Mackenson Cineus débute ses études en Gestion des Affaires à l\'<a href="https://www.ueh.edu.ht" target="_blank" style="color:#C9A84C;"><b>INAGHEI</b></a> '
            '(Université d\'État d\'Haïti), en parallèle d\'un double cursus en Sciences Économiques (FDSE). '
            'Dès 2016, il s\'engage à <b>ACTIVEH</b> comme co-leader du pôle Partenariat & Développement, '
            'puis Président du <b>Groupe Économie & Finance</b> — organisant des conférences, '
            'lançant le projet « Goud la se Pa m » (sensibilisation monétaire, musée BRH, 3 émissions radio).',

            'Il devient <b>Directeur des Opérations puis Directeur Financier de '
            '<a href="https://www.hultprize.org" target="_blank" style="color:#C9A84C;">Hult Prize Haïti</a></b>, '
            'représente Haïti au Hult Prize international 2018 avec le projet <b>WEISS</b> '
            '(transformation des déchets en biogaz & fertilisant), et cofonde <b>ANGAJMAN</b> '
            'comme Directeur Exécutif. En parallèle, il milite au '
            '<b>PHJEA</b> (Parlement Haïtien Jeunesse Eau) et remporte le '
            '<b>Hackathon Unleash</b> (Objectifs de Développement Durable) avec une équipe internationale.',
        ],
        "tags": ["ACTIVEH", "Hult Prize 2018 · WEISS", "ANGAJMAN", "PHJEA · Eau", "Unleash · ODD"],
    },
    {
        "phase": "phase2",
        "label": "Phase 2 · 2019–2024 · France",
        "title": "Formation Finance & Engagement international",
        "years": "Université du Mans · Erasmus Expertise · Podcast INCLUTECH · MBA ESLSCA",
        "text": [
            'Arrivé en France, Mackenson obtient une <b>Licence Banque, Finance & Assurance</b> '
            'à l\'<a href="https://www.univ-lemans.fr" target="_blank" style="color:#C9A84C;"><b>Université du Mans</b></a> '
            '(2022–2024), puis intègre le <b>MBA 2 Finance de Marché</b> (parcours anglais) '
            'à l\'<a href="https://www.eslsca.fr" target="_blank" style="color:#C9A84C;"><b>ESLSCA Business School Paris</b></a> '
            '(2024–2026). Son mémoire de 80 pages porte sur l\'impact des réglementations ESG sur la performance des portefeuilles.',

            'Il rejoint <a href="https://erasmus-expertise.eu" target="_blank" style="color:#C9A84C;"><b>Erasmus Expertise</b></a> '
            'comme tuteur (7 étudiants accompagnés) avant d\'être nommé <b>Board Member</b> en 2023. '
            'Il lance le podcast '
            '<a href="https://open.spotify.com/show/5XvFdWYwhHWY3EguIhhf69" target="_blank" style="color:#C9A84C;"><b>INCLUTECH</b></a> '
            '(Spotify & Apple, bi-mensuel) sur l\'inclusion financière. '
            'En octobre 2023, son équipe remporte la <b>1ère place du Hackathon Fintech Générations</b> '
            '(France FinTech / Société Générale / Treezor) avec le projet <b>Victoria</b> (financement DPE).',
        ],
        "tags": ["Licence BFA · Le Mans", "MBA ESLSCA · Paris", "Erasmus Expertise", "Podcast INCLUTECH", "Hackathon 2023 🥇"],
    },
    {
        "phase": "phase3",
        "label": "Phase 3 · 2024–2025 · HiPay SAS",
        "title": "Chargé de conformité — Fintech de paiement agréée ACPR",
        "years": "HiPay SAS · Levallois-Perret · LCB-FT · Pays sensibles · PPE · Tracfin",
        "text": [
            '<a href="https://hipay.com" target="_blank" style="color:#C9A84C;"><b>HiPay SAS</b></a> '
            '(fintech de paiement omnicanal, Euronext Growth ALHYP, agréée ACPR). '
            'Missions : analyse risques pays GAFI sur <b>pays sensibles et complexes</b> '
            '(sanctions internationales, jeux en ligne, marketplaces) ; '
            'analyses KYC/PPE/SOE ; réponses aux réquisitions '
            '<a href="https://www.economie.gouv.fr/tracfin" target="_blank" style="color:#C9A84C;"><b>Tracfin</b></a> '
            'et judiciaires ; Due Diligence établissements financiers systémiques ; '
            'mise à jour de la cartographie des risques.',
        ],
        "tags": ["HiPay · ALHYP", "LCB-FT · PPE/SOE", "Pays sensibles · GAFI", "Tracfin", "Due Diligence"],
    },
    {
        "phase": "phase4",
        "label": "Phase 4 · 2025–Présent · Banque DELUBAC & CIE",
        "title": "Analyste LCB-FT — Corporate, CIB, Asset Management, PSP",
        "years": "Banque DELUBAC & CIE · Paris · Alertes N2 · ER · KYC · DS/SAR · Réquisitions",
        "text": [
            '<a href="https://www.delubac.com" target="_blank" style="color:#C9A84C;"><b>Banque Delubac & Cie</b></a> '
            '(fondée 1924, banque indépendante de plein exercice). '
            'Analyste LCB-FT en alternance dans le cadre du MBA 2 Finance de Marché (ESLSCA). '
            'Missions : traitement <b>alertes de niveau 2</b> sur clientèle corporate, CIB, '
            'asset management, banque privée ; <b>Examens Renforcés (ER)</b> ; '
            'analyse KYC et identification des bénéficiaires effectifs ; '
            '<b>déclarations de soupçon (DS)</b> et règles d\'exclusion sur transactions récurrentes ; '
            'analyse des états financiers clients ; réquisitions judiciaires.',

            'Son dashboard Excel de suivi couvre 205 analyses pour un volume de 24,9 M€ examinés, '
            'avec typologies : blanchiment, travail dissimulé, abus de biens sociaux, '
            'financement du terrorisme. Double certification '
            '<a href="https://acpr.banque-france.fr" target="_blank" style="color:#C9A84C;"><b>ACPR</b></a> '
            '(AMF) et CAMS (Sanctions & Embargos).',
        ],
        "tags": ["Banque DELUBAC", "Alertes N2 · ER", "KYC corporate/CIB/PSP", "DS/SAR", "ACPR · GAFI", "AMF · CAMS"],
    },
]

# ── BIO COURTE (2-3 paragraphes segmentés) ────────────────────────────────────
BIO_SHORT = {
    "formation": (
        "Titulaire d'une Licence en Banque, Finance et Assurance (Université du Mans, 2022–2024) "
        "et actuellement en MBA 2 Finance de Marché (parcours anglais) à l'ESLSCA Business School "
        "Paris (2024–2026), Mackenson CINÉUS a bâti une solide formation académique en finance de "
        "marché et conformité réglementaire. Son mémoire de recherche de 80 pages, rédigé en anglais, "
        "porte sur l'impact des facteurs ESG et des contraintes réglementaires sur la gestion de "
        "portefeuille — témoignant de sa capacité à traiter des problématiques complexes à "
        "l'intersection de la réglementation et de la performance financière. Certifié AMF et CAMS "
        "(Sanctions & Embargos), il prépare également le CFA Level I."
    ),
    "experience": (
        "Fort d'une double expérience opérationnelle en conformité LCB-FT, Mackenson a exercé "
        "comme Chargé de Conformité chez HiPay SAS (fintech de paiement agréée ACPR, 2024–2025) "
        "— analyse des risques pays GAFI, KYC/PPE/SOE sur pays sensibles et complexes, "
        "réponses aux réquisitions Tracfin — puis comme Analyste LCB-FT en alternance à la "
        "Banque Delubac & Cie (2025–présent) — traitement des alertes de niveau 2, Examens "
        "Renforcés, déclarations de soupçon sur clientèle corporate, CIB et asset management. "
        "Lauréat du 1er Prix Hackathon Fintech Générations 2023 (France FinTech / Société Générale, "
        "Projet Victoria), il illustre sa capacité à allier rigueur réglementaire et innovation."
    ),
    "engagement": (
        "Parallèlement à son parcours professionnel, Mackenson s'est toujours engagé dans des "
        "initiatives à fort impact social et environnemental. En Haïti, il a cofondé ANGAJMAN "
        "(projets citoyens dans 10 départements), dirigé le Hult Prize Haïti (Projet WEISS : "
        "biogaz & fertilisant organique), milité au PHJEA pour l'eau potable et remporté le "
        "Hackathon Unleash sur les Objectifs de Développement Durable avec une équipe internationale. "
        "En France, il est Board Member d'Erasmus Expertise (ONG internationale), crée et anime "
        "le podcast INCLUTECH (Spotify & Apple Podcasts) sur l'inclusion financière et la fintech, "
        "et continue de construire un réseau franco-international dans l'écosystème de la conformité "
        "et des marchés financiers."
    ),
}

PLATFORMS = [
    {"icon":"💼","name":"LinkedIn","desc":"Profil professionnel · réseau finance/compliance.","url":"https://www.linkedin.com/in/mackenson-cineus","label":"LinkedIn →"},
    {"icon":"💳","name":"HiPay","desc":"Fintech paiement ACPR agréée — 2024–2025.","url":"https://hipay.com","label":"HiPay →"},
    {"icon":"🏦","name":"Banque Delubac","desc":"Banque indépendante — Analyste LCB-FT 2025.","url":"https://www.delubac.com","label":"Delubac →"},
    {"icon":"🚀","name":"France FinTech","desc":"Lauréat Hackathon 2023 · Membre actif.","url":"https://francefintech.org","label":"France FinTech →"},
    {"icon":"🏆","name":"Hult Prize","desc":"Représentant Haïti 2018 · Projet WEISS.","url":"https://www.hultprize.org","label":"Hult Prize →"},
    {"icon":"🎓","name":"ESLSCA Paris","desc":"MBA 2 Finance de marché 2024–2026.","url":"https://www.eslsca.fr","label":"ESLSCA →"},
    {"icon":"🏫","name":"Université du Mans","desc":"Licence BFA 2022–2024.","url":"https://www.univ-lemans.fr","label":"Univ. du Mans →"},
    {"icon":"🎙","name":"Podcast INCLUTECH","desc":"Finance & Inclusion · Spotify & Apple.","url":"https://open.spotify.com/show/5XvFdWYwhHWY3EguIhhf69","label":"Écouter →"},
    {"icon":"🌍","name":"Erasmus Expertise","desc":"Board Member · ONG internationale.","url":"https://erasmus-expertise.eu","label":"Erasmus →"},
    {"icon":"📸","name":"Instagram","desc":"@mackenson_cineus","url":"https://instagram.com/mackenson_cineus","label":"Instagram →"},
]

REGULATORS = [
    ("🇫🇷","ACPR","Régulateur bancaire & paiement · LCB-FT","https://acpr.banque-france.fr"),
    ("🌐","FATF / GAFI","Normes mondiales anti-blanchiment","https://www.fatf-gafi.org"),
    ("💰","Tracfin","Cellule de Renseignement Financier","https://www.economie.gouv.fr/tracfin"),
    ("📊","AMF","Autorité des Marchés Financiers","https://www.amf-france.org"),
]

_EXTRA_CSS = """
<style>
  .bio-period.phase5::before { background:linear-gradient(180deg,#C9A84C,#E8C97A); }
  .phase5 .bio-phase-label { color:#C9A84C; }
</style>
"""

NAVY_C    = rl_colors.HexColor('#0D1B2A')
GOLD_C    = rl_colors.HexColor('#C9A84C')
GRAY_D_C  = rl_colors.HexColor('#333333')
GRAY_M_C  = rl_colors.HexColor('#555555')
GRAY_L_C  = rl_colors.HexColor('#999999')
CREAM_C   = rl_colors.HexColor('#F8F6F2')
ACCENT_C  = rl_colors.HexColor('#4A9EFF')
NAVY_M_C  = rl_colors.HexColor('#1A2E42')

def _s(name, **kw):
    d = dict(fontName='Helvetica', fontSize=9, textColor=GRAY_D_C, leading=13, spaceAfter=3)
    d.update(kw)
    return ParagraphStyle(name, **d)

def _hr(color=GOLD_C, t=1.0, before=2, after=5):
    return HRFlowable(width="100%", thickness=t, color=color, spaceBefore=before, spaceAfter=after)

def _sec(label, size=7, letter=1.5):
    return [
        Paragraph(f'<font color="#C9A84C"><b>{label.upper()}</b></font>',
                  _s('sh', fontSize=size, letterSpacing=letter, spaceAfter=1, spaceBefore=7)),
        HRFlowable(width="100%", thickness=0.7, color=GOLD_C, spaceAfter=3, spaceBefore=0),
    ]

def _page_frame(canvas, doc):
    """Shared header/footer for all PDFs."""
    W, H = A4
    canvas.saveState()
    canvas.setFillColor(NAVY_C); canvas.rect(0, H-1.55*cm, W, 1.55*cm, fill=1, stroke=0)
    canvas.setFillColor(GOLD_C); canvas.rect(0, H-1.55*cm, W, 0.09*cm, fill=1, stroke=0)
    canvas.setFillColor(rl_colors.white); canvas.setFont("Helvetica-Bold", 9)
    canvas.drawString(1.8*cm, H-1.0*cm, "CINÉUS Mackenson")
    canvas.setFillColor(GOLD_C); canvas.setFont("Helvetica", 7.5)
    canvas.drawString(5.8*cm, H-1.0*cm, "Finance · Compliance · Fintech · Paris")
    canvas.setFillColor(NAVY_C); canvas.rect(0, 0, W, 0.85*cm, fill=1, stroke=0)
    canvas.setFillColor(GOLD_C); canvas.rect(0, 0.85*cm, W, 0.07*cm, fill=1, stroke=0)
    canvas.setFillColor(rl_colors.HexColor('#888888')); canvas.setFont("Helvetica", 7)
    canvas.drawString(1.8*cm, 0.3*cm,
        "linkedin.com/in/mackenson-cineus  ·  Podcast INCLUTECH  ·  @mackenson_cineus")
    canvas.setFillColor(GOLD_C)
    canvas.drawRightString(W-1.8*cm, 0.3*cm, f"Page {doc.page}")
    canvas.restoreState()


# ─────────────────────────────────────────────────────────────────────────────
# CV — 1 PAGE BICOLONNE FINANCE
# ─────────────────────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────────────────────
# CV — 1 PAGE · FORMAT FINANCE EUROPÉEN STANDARD
# Mise en page entièrement dessinée sur canvas pour un contrôle parfait
# ─────────────────────────────────────────────────────────────────────────────




# ═══════════════════════════════════════════════════════════════════════════════
# PDF HELPERS
# ═══════════════════════════════════════════════════════════════════════════════
def _sw(s, fn, fs):
    from reportlab.pdfbase.pdfmetrics import stringWidth
    return stringWidth(str(s), fn, fs)

def _wrap_canvas(c, txt, x, y, mw, fn, fs, lh, col=(0,0,0), ind=0):
    c.setFillColorRGB(*col); c.setFont(fn, fs)
    words = str(txt).split(); cur = ""
    for w in words:
        test = (cur+" "+w).strip()
        if _sw(test, fn, fs) <= mw - ind: cur = test
        else:
            if cur: c.drawString(x+ind, y, cur); y -= lh; ind = 0
            cur = w
    if cur: c.drawString(x+ind, y, cur); y -= lh
    return y

def _line_h(c, y, x1, x2, lw=0.5, col=(0,0,0)):
    c.setStrokeColorRGB(*col); c.setLineWidth(lw)
    c.line(x1, y, x2, y)

# ── SHARED BULLET CONTENT ────────────────────────────────────────────────────
def _get_cv_data(lang="fr"):
    ss = st.session_state
    exps = ss.get("edit_exp", [])
    edus = ss.get("edit_edu", [])
    skills = ss.get("edit_skills", {})
    certifs = ss.get("edit_certifs", [])
    tech = ss.get("edit_tech", "Excel – VBA – Python – Bloomberg – Looker – Jura")
    ints = ss.get("edit_interests", "Sécurité financière – Modèles financiers – Trading NASDAQ")
    dists = ss.get("edit_distinctions", [])
    memoir = ss.get("edit_memoir_title",
                    "Impacts of ESG Factors and Regulatory Constraints on Portfolio Management")
    name = ss.get("edit_name", "Mackenson CINÉUS")
    title = ss.get("edit_title", "Analyste LCB-FT | Compliance Officer")
    loc = ss.get("edit_location", "Île-de-France, France")
    email = ss.get("edit_email", "mackenson.cineus@email.com")
    phone = ss.get("edit_phone", "+33 6 XX XX XX XX")
    lk = ss.get("edit_linkedin", "linkedin.com/in/mackenson-cineus")
    return dict(exps=exps, edus=edus, skills=skills, certifs=certifs, tech=tech,
                ints=ints, dists=dists, memoir=memoir, name=name, title=title,
                loc=loc, email=email, phone=phone, lk=lk, lang=lang)


# ═══════════════════════════════════════════════════════════════════════════════
# CV PDF — 1 PAGE COMPLÈTE, NOIR/GRIS, POLICE 9.5, EN-TÊTE CENTRÉ
# ═══════════════════════════════════════════════════════════════════════════════
def generate_cv_pdf(poste="", entreprise="", secteur="", contexte="", lang="fr") -> bytes:
    from reportlab.pdfgen import canvas as rl_canvas
    buf = io.BytesIO()
    W, H = A4
    d = _get_cv_data(lang)
    target = poste if poste else d["title"]

    ML, MR, MT, MB = 44, 44, 36, 24
    TW = W - ML - MR

    BLK  = (0.0,  0.0,  0.0)
    DGR  = (0.20, 0.20, 0.20)
    MGR  = (0.40, 0.40, 0.40)
    LGR  = (0.60, 0.60, 0.60)
    VGR  = (0.82, 0.82, 0.82)

    FS = 9.5       # base font size (all body text)
    LH = 11.2      # line height

    c = rl_canvas.Canvas(buf, pagesize=A4)

    def f(fn, fs): c.setFont(fn, fs)
    def sf(*rgb): c.setFillColorRGB(*rgb)
    def line(y, w=0.6, col=BLK, x1=ML, x2=None):
        if x2 is None: x2 = ML+TW
        c.setStrokeColorRGB(*col); c.setLineWidth(w); c.line(x1, y, x2, y)
    def wrap(txt, x, y, mw, fn, fs, lh, col=BLK, ind=0):
        return _wrap_canvas(c, txt, x, y, mw, fn, fs, lh, col, ind)

    def sec(label, y, sp_before=5):
        y -= sp_before
        lbl = label.upper() if lang=="fr" else label.upper()
        f("Times-Bold", 9.8); sf(*BLK)
        c.drawString(ML, y, lbl); y -= 2
        line(y, 0.9, BLK); return y - 8

    def bul(txt, y, fs=FS, lh=LH, col=DGR, x=ML, mw=None):
        if mw is None: mw = TW
        f("Times-Roman", fs); sf(*col)
        c.drawString(x+7, y, "•")
        return wrap(txt, x+18, y, mw-18, "Times-Roman", fs, lh, col)

    # ── EN-TÊTE CENTRÉ ─────────────────────────────────────────────
    y = H - MT
    f("Times-Bold", 20); sf(*BLK)
    nw = _sw(d["name"].upper(), "Times-Bold", 20)
    c.drawString((W-nw)/2, y, d["name"].upper()); y -= 14

    f("Times-Italic", FS); sf(*DGR)
    tw = _sw(target, "Times-Italic", FS)
    c.drawString((W-tw)/2, y, target); y -= 12

    line(y, 0.35, VGR); y -= 8

    f("Times-Roman", 9); sf(*MGR)
    coords_parts = [d["loc"], d["phone"], d["email"]]
    row1 = "  ·  ".join(coords_parts)
    row2 = f"{d['lk']}  ·  Podcast INCLUTECH"
    r1w = _sw(row1, "Times-Roman", 9)
    r2w = _sw(row2, "Times-Roman", 9)
    c.drawString((W-r1w)/2, y, row1); y -= 10
    c.drawString((W-r2w)/2, y, row2); y -= 10
    line(y, 0.9, BLK); y -= 2; line(y, 0.3, VGR); y -= 8

    # ── PROFIL ─────────────────────────────────────────────────────
    y = sec("Profil" if lang=="fr" else "Profile", y, 2)
    if lang == "fr":
        profil = (
            "Analyste LCB-FT en alternance chez Banque Delubac & Cie sur clientèle corporate, CIB, "
            "asset management et banque privée. Expérience Chargé de Conformité chez HiPay SAS "
            "(fintech agréée ACPR) : risques pays GAFI, KYC/PPE/SOE, pays sensibles et complexes, "
            "Tracfin, Due Diligence. MBA 2 Finance de Marché (ESLSCA Paris, parcours anglais). "
            "Certifié AMF, CAMS (Sanctions & Embargos), candidat CFA Level I. "
            "Lauréat Hackathon Fintech Générations 2023 — France FinTech / Société Générale."
        )
    else:
        profil = d.get("summary", (
            "Work-study AML/CFT Analyst at Banque Delubac & Cie — Corporate, CIB, Asset Management clients. "
            "Fintech Compliance experience at HiPay SAS (ACPR-licensed): FATF risk countries, KYC/PEP/SOE, "
            "high-risk jurisdictions, Tracfin, Due Diligence. MBA in Market Finance (ESLSCA Paris, English track). "
            "AMF Certified, CAMS (Sanctions & Embargoes), CFA Level I Candidate. "
            "1st Prize Hackathon Fintech Générations 2023 — France FinTech / Société Générale."
        ))
    y = wrap(profil, ML, y, TW, "Times-Roman", FS, LH, DGR); y -= 5

    # ── EXPÉRIENCES ────────────────────────────────────────────────
    sec_label = "Expériences Professionnelles" if lang=="fr" else "Professional Experience"
    y = sec(sec_label, y)

    PRO = {"Banque","Fintech","Compétition","ONG",""}
    pro_exps = [e for e in d["exps"]
                if e.get("org_type","") in PRO
                and "ACTIVEH" not in e.get("org","")
                and "ANGAJMAN" not in e.get("org","")
                and "Hult Prize Haïti" not in e.get("org","")][:4]

    for exp in pro_exps:
        if y < MB+55: break
        org_parts = exp.get("org","").split("·")
        org_name = org_parts[0].strip()
        org_loc  = org_parts[-1].strip() if len(org_parts)>1 else ("Paris, France" if lang=="fr" else "Paris, France")
        role  = exp.get("role_en" if lang=="en" else "role", exp.get("role",""))
        period = exp.get("period_en" if lang=="en" else "period", exp.get("period",""))
        bullets = exp.get("bullets_en" if lang=="en" else "bullets_fr",
                          exp.get("bullets_fr", exp.get("bullets",[])))

        f("Times-Bold", FS); sf(*BLK)
        c.drawString(ML, y, org_name)
        f("Times-Roman", 9); sf(*LGR)
        c.drawRightString(ML+TW, y, org_loc)
        y -= 11

        f("Times-BoldItalic", FS-0.3); sf(*DGR)
        c.drawString(ML, y, role)
        f("Times-Italic", 9); sf(*LGR)
        c.drawRightString(ML+TW, y, period)
        y -= 11

        for b in bullets[:5]:
            if y < MB+30: break
            y = bul(b, y)

        y -= 3; line(y, 0.28, VGR); y -= 4

    # ── MÉMOIRE ────────────────────────────────────────────────────
    if y > MB+78:
        mem_lbl = "Mémoire de Fin d'Études" if lang=="fr" else "Master's Thesis"
        y = sec(mem_lbl, y)
        memoir_full = d["memoir"]
        f("Times-BoldItalic", FS); sf(*BLK)
        lbl = ("Titre : " if lang=="fr" else "Title: ")
        c.drawString(ML, y, lbl)
        lw = _sw(lbl,"Times-BoldItalic",FS)
        f("Times-Italic", FS); sf(*DGR)
        y = wrap(f"« {memoir_full} »", ML+lw, y, TW-lw, "Times-Italic", FS, LH, DGR)
        f("Times-Roman", FS-0.5); sf(*MGR)
        line1 = ("80 pages · Rédigé en anglais · MBA 2 Finance de Marché, ESLSCA Paris"
                 if lang=="fr" else "80 pages · Written in English · MBA in Market Finance, ESLSCA Paris")
        c.drawString(ML, y, line1); y -= LH
        line2 = ("Performance des portefeuilles sous contrainte réglementaire ESG/non-ESG"
                 if lang=="fr" else "Portfolio performance under ESG/non-ESG regulatory constraints")
        c.drawString(ML, y, line2); y -= LH+3

    # ── FORMATION ──────────────────────────────────────────────────
    if y > MB+65:
        form_lbl = "Formations" if lang=="fr" else "Education"
        y = sec(form_lbl, y)

        EXCL = ["INAGHEI","Haïti","Haiti"]
        edus_f = [e for e in d["edus"] if not any(x in e.get("school","") for x in EXCL)][:2]

        for edu in edus_f:
            if y < MB+40: break
            s_parts = edu.get("school","").split("·")
            sname = s_parts[0].strip()
            sloc  = s_parts[-1].strip() if len(s_parts)>1 else "France"
            deg   = edu.get("deg_en" if lang=="en" else "deg", edu.get("deg",""))
            yr    = edu.get("yr_en"  if lang=="en" else "yr",  edu.get("yr",""))
            det   = edu.get("det_en" if lang=="en" else "det", edu.get("det",""))

            f("Times-Bold", FS); sf(*BLK)
            c.drawString(ML, y, f"{sname}  |  {deg}")
            f("Times-Roman", 9); sf(*LGR)
            c.drawRightString(ML+TW, y, sloc)
            y -= 11
            f("Times-Italic", 9); sf(*LGR)
            c.drawRightString(ML+TW, y, yr)
            clbl = ("Cours principaux : " if lang=="fr" else "Main courses: ")
            f("Times-Roman", FS); sf(*BLK)
            c.drawString(ML, y, clbl)
            cw = _sw(clbl,"Times-Roman",FS)
            c.setStrokeColorRGB(*BLK); c.setLineWidth(0.25)
            c.line(ML, y-1, ML+cw, y-1)
            f("Times-Roman", FS); sf(*MGR)
            y = wrap(det, ML+cw, y, TW-cw, "Times-Roman", FS, LH, MGR)
            y -= 4

    # ── CERTIFICATIONS, COMPÉTENCES, LANGUES ──────────────────────
    if y > MB+5:
        cert_lbl = ("Certifications, Compétences, Langues & Distinctions"
                    if lang=="fr" else "Certifications, Skills, Languages & Awards")
        y = sec(cert_lbl, y)

        certif_str = " – ".join(d["certifs"]) if d["certifs"] else "AMF – CAMS – CFA Level I"
        dist_str = " | ".join(f"{ic} {t}" for ic,t,_ in d["dists"][:3])

        rows_fr = [
            ("Certifications",        f": {certif_str}"),
            ("Outils & Informatique", f": {d['tech']}"),
            ("Langues",               ": Français (Natif) – Anglais (Avancé – C1)"),
            ("Intérêts",              f": {d['ints']}"),
            ("Distinctions",          f": {dist_str}"),
        ]
        rows_en = [
            ("Certifications",        f": {certif_str}"),
            ("Tools & IT",            f": {d['tech']}"),
            ("Languages",             ": French (Native) – English (Advanced – C1)"),
            ("Interests",             f": {d['ints']}"),
            ("Awards",                f": {dist_str}"),
        ]
        rows = rows_fr if lang=="fr" else rows_en

        for lbl_r, val in rows:
            if y < MB: break
            f("Times-Bold", FS); sf(*BLK)
            lw = _sw(lbl_r,"Times-Bold",FS)
            c.drawString(ML+6, y, lbl_r)
            c.setStrokeColorRGB(*BLK); c.setLineWidth(0.25)
            c.line(ML+6, y-1, ML+6+lw, y-1)
            f("Times-Roman", FS); sf(*MGR)
            y = wrap(val, ML+6+lw, y, TW-6-lw, "Times-Roman", FS, LH, MGR)

    # ── FOOTER ─────────────────────────────────────────────────────
    c.setStrokeColorRGB(*VGR); c.setLineWidth(0.3)
    c.line(ML, MB+5, ML+TW, MB+5)
    f("Helvetica", 6); sf(*LGR)
    foot = f"{d['name']}  ·  {d['email']}  ·  {d['lk']}"
    fw = _sw(foot,"Helvetica",6)
    c.drawString((W-fw)/2, MB-4, foot)

    c.save(); buf.seek(0); return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# CV WORD — .docx export (FR or EN)
# ═══════════════════════════════════════════════════════════════════════════════

def generate_cv_docx(poste="", lang="fr") -> bytes:
    """
    CV DOCX matching EXACTLY the uploaded model:
    - Header table: Name 18pt bold centered, title, contacts
    - Section headers: 10pt bold, bottom border, uppercase
    - Body text: 9.5pt
    - Bullets: List Bullet style
    - Org+date pattern: bold org | italic loc (tab right), date right-aligned
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    import streamlit as _st
    ss = _st.session_state
    is_fr = (lang == "fr")

    # Live data from session_state
    name     = ss.get("edit_name",    "Mackenson CINÉUS")
    phone    = ss.get("edit_phone",   "+33 6 XX XX XX XX")
    email    = ss.get("edit_email",   "mackenson.cineus@email.com")
    linkedin = ss.get("edit_linkedin","linkedin.com/in/mackenson-cineus")
    title_cv = poste if poste else ss.get("edit_title","Analyste LCB-FT | Compliance Officer")
    memoir   = ss.get("edit_memoir_title","Impacts of ESG Factors and Regulatory Constraints on Portfolio Management")
    certifs  = ss.get("edit_certifs", ["AMF","CAMS (Sanctions & Embargos)","CFA Level I (en cours)"])
    tech     = ss.get("edit_tech",    "Excel avancé – VBA – Python – Pack Microsoft Office – Looker – Jura")
    interests= ss.get("edit_interests","Sécurité financière – Création de modèles financiers – Programmation")
    dists    = ss.get("edit_distinctions",[])
    exps     = ss.get("edit_exp",    [])
    edus     = ss.get("edit_edu",    [])

    profil_fr = ss.get("edit_summary_fr",
        "Analyste LCB-FT en alternance chez Banque Delubac & Cie avec expertise en surveillance des "
        "transactions sur clientèle corporate, CIB, asset management et banque privée. Expérience en "
        "conformité Fintech (HiPay SAS) axée sur les risques pays GAFI, KYC/PPE/SOE et juridictions à "
        "haut risque. MBA 2 Finance de Marché (ESLSCA Paris, parcours anglais). Certifié AMF, CAMS "
        "(Sanctions & Embargos), candidat CFA Level I.")
    profil_en = ss.get("edit_summary_en",
        "Work-study AML/CFT Analyst at Banque Delubac & Cie with expertise in transaction monitoring for "
        "Corporate, CIB, Asset Management and Private Banking clients. Proven experience in Fintech "
        "Compliance (HiPay SAS) focusing on FATF risk countries, KYC/PEP/SOE, and high-risk jurisdictions. "
        "MBA in Market Finance at ESLSCA Paris (English Track). AMF Certified, CAMS (Sanctions & Embargoes), "
        "CFA Level I Candidate.")

    doc = Document()

    # Margins (match model: 0.75in L/R, 0.6in T/B)
    for sec in doc.sections:
        sec.top_margin    = Inches(0.6)
        sec.bottom_margin = Inches(0.6)
        sec.left_margin   = Inches(0.75)
        sec.right_margin  = Inches(0.75)

    # ── HELPERS ────────────────────────────────────────────────────────────
    def bottom_border(p):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        b = OxmlElement('w:bottom')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'12')
        b.set(qn('w:space'),'1');    b.set(qn('w:color'),'000000')
        pBdr.append(b); pPr.append(pBdr)

    def sec_hdr(label, sb=8):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(label); r.bold = True
        r.font.size = Pt(10); r.font.color.rgb = RGBColor(0,0,0)
        bottom_border(p); return p

    def org_line(org_text, loc_text):
        """Bold org | italic loc with right-tab."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(0)
        r1 = p.add_run(org_text); r1.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run("\t" + loc_text)
        r2.italic = True; r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor(40,40,40)
        # Right tab at 9360 twips (6.5in printable width)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab  = OxmlElement('w:tab')
        tab.set(qn('w:val'),'right'); tab.set(qn('w:pos'),'9360')
        tabs.append(tab); pPr.append(tabs)

    def date_line(date_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(date_text); r.italic = True
        r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(40,40,40)

    def bullet(text, size=9.5):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(text); r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0,0,0)

    def cours_line(label_bold, cours_text, size=9.5):
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label_bold); r1.bold = True; r1.font.size = Pt(size)
        r2 = p.add_run(cours_text); r2.font.size = Pt(size)
        r2.font.color.rgb = RGBColor(0,0,0)

    def certif_line(label_bold, value, size=9.5):
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label_bold); r1.bold = True; r1.font.size = Pt(size)
        r2 = p.add_run(value);      r2.font.size = Pt(size)
        r2.font.color.rgb = RGBColor(0,0,0)

    # ── HEADER TABLE ────────────────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=1)
    # Remove all borders
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl._tbl.insert(0, tblPr)
    tblB = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'none'); tblB.append(el)
    tblPr.append(tblB)

    cell = tbl.cell(0, 0)
    cell.paragraphs[0].clear()

    # Name — 18pt bold centered black
    pn = cell.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn.paragraph_format.space_before = Pt(0); pn.paragraph_format.space_after = Pt(2)
    rn = pn.add_run(name.upper())
    rn.bold = True; rn.font.size = Pt(18); rn.font.color.rgb = RGBColor(0,0,0)

    # Title
    pt = cell.add_paragraph(); pt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pt.paragraph_format.space_before = Pt(0); pt.paragraph_format.space_after = Pt(2)
    rt = pt.add_run(title_cv)
    rt.font.size = Pt(10.5); rt.font.color.rgb = RGBColor(40,40,40)

    # Contacts
    pc = cell.add_paragraph(); pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.space_before = Pt(0); pc.paragraph_format.space_after = Pt(4)
    rc = pc.add_run(f"{phone}  |  {linkedin}  |  {email}")
    rc.font.size = Pt(10); rc.font.color.rgb = RGBColor(40,40,40)

    # Spacing after header
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(0); sp.paragraph_format.space_after = Pt(2)

    # ── PROFIL ──────────────────────────────────────────────────────────────
    sec_hdr("PROFIL" if is_fr else "PROFILE", sb=2)
    p_pr = doc.add_paragraph()
    p_pr.paragraph_format.space_before = Pt(4); p_pr.paragraph_format.space_after = Pt(1)
    rp = p_pr.add_run(profil_fr if is_fr else profil_en)
    rp.font.size = Pt(9.5); rp.font.color.rgb = RGBColor(0,0,0)

    # ── EXPÉRIENCES ─────────────────────────────────────────────────────────
    sec_hdr("EXPÉRIENCES PROFESSIONNELLES" if is_fr else "PROFESSIONAL EXPERIENCE")

    # Filter professional exps
    PRO = {"Banque","Fintech","Compétition","ONG",""}
    pro_exps = [e for e in exps
                if e.get("org_type","") in PRO
                and "ACTIVEH" not in e.get("org","")
                and "ANGAJMAN" not in e.get("org","")]

    for exp in pro_exps:
        org_raw = exp.get("org","")
        org_parts = org_raw.split("·")
        org_name  = org_parts[0].strip()
        org_loc   = org_parts[-1].strip() if len(org_parts)>1 else "France"
        role  = exp.get("role_en" if not is_fr else "role", exp.get("role",""))
        period= exp.get("period_en" if not is_fr else "period", exp.get("period",""))
        buls  = exp.get("bullets_en" if not is_fr else "bullets_fr",
                        exp.get("bullets_fr", exp.get("bullets",[])))
        # Format org line: "ORG NAME | ROLE" then loc
        org_line(f"{org_name} | {role}", org_loc)
        date_line(period)
        for b in buls[:5]:
            bullet(b)

    # ── FORMATIONS ──────────────────────────────────────────────────────────
    sec_hdr("FORMATIONS" if is_fr else "EDUCATION")
    EXCL = ["INAGHEI","Haïti","Haiti"]
    edus_f = [e for e in edus if not any(x in e.get("school","") for x in EXCL)]

    for edu in edus_f:
        s_parts = edu.get("school","").split("·")
        sname = s_parts[0].strip()
        sloc  = s_parts[-1].strip() if len(s_parts)>1 else "France"
        deg   = edu.get("deg_en" if not is_fr else "deg", edu.get("deg",""))
        yr    = edu.get("yr_en"  if not is_fr else "yr",  edu.get("yr",""))
        det   = edu.get("det_en" if not is_fr else "det", edu.get("det",""))
        org_line(f"{sname} — {deg}", sloc)
        date_line(yr)
        label = "Cours principaux\u00a0: " if is_fr else "Main courses: "
        cours_line(label, det)

    # ── MÉMOIRE ─────────────────────────────────────────────────────────────
    sec_hdr("MÉMOIRE DE FIN D'ÉTUDES" if is_fr else "MASTER'S THESIS")
    pm = doc.add_paragraph()
    pm.paragraph_format.space_before = Pt(4); pm.paragraph_format.space_after = Pt(1)
    lbl_m = "Titre\u00a0: " if is_fr else "Title: "
    rm1 = pm.add_run(lbl_m); rm1.bold = True; rm1.font.size = Pt(9.5)
    memoir_desc = (
        f"« {memoir} » — 80 pages · " +
        ("Rédigé en anglais · MBA 2 Finance de Marché · ESLSCA Paris" if is_fr else
         "Written in English · MBA in Market Finance · ESLSCA Paris")
    )
    rm2 = pm.add_run(memoir_desc); rm2.italic = True; rm2.font.size = Pt(9.5)
    rm2.font.color.rgb = RGBColor(0,0,0)

    # ── CERTIFICATIONS ───────────────────────────────────────────────────────
    cert_lbl_hdr = ("CERTIFICATIONS, COMPÉTENCES, LANGUES & DISTINCTIONS" if is_fr
                    else "CERTIFICATIONS, SKILLS, LANGUAGES & AWARDS")
    sec_hdr(cert_lbl_hdr)

    certif_str = " – ".join(certifs) if certifs else "AMF – CAMS – CFA Level I"
    dist_str = " | ".join(f"{ic} {t}" for ic,t,_ in dists[:4]) if dists else \
               "Hackathon Fintech Générations 2023 | Hult Prize 2018 | Hackathon Unleash"

    rows_fr = [
        ("Certifications\u00a0: ", certif_str),
        ("Outils & Informatique\u00a0: ", tech),
        ("Langues\u00a0: ", "Français (Natif) – Anglais (Avancé – C1)"),
        ("Intérêts\u00a0: ", interests),
        ("Distinctions\u00a0: ", dist_str),
    ]
    rows_en = [
        ("Certifications: ", certif_str),
        ("Tools & IT: ", tech),
        ("Languages: ", "English (Advanced C1) – French (Native) – Haitian Creole (Native)"),
        ("Interests: ", interests),
        ("Awards: ", dist_str),
    ]
    for lbl, val in (rows_fr if is_fr else rows_en):
        certif_line(lbl, val)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# JOB SEARCH — REAL SCRAPER using Claude API web_search
# ═══════════════════════════════════════════════════════════════════════════════
def _search_real_jobs(api_key: str, keywords: list, location: str = "France") -> list:
    """
    Use Claude API with web_search tool to get REAL current CDI job listings
    in financial security / AML compliance.
    Returns list of dicts: {title, company, location, url, date, description}
    """
    import requests as _req, json

    if not api_key:
        return []

    # Build a comprehensive search query
    kw_str = " OR ".join(f'"{k}"' for k in keywords[:5])
    prompt = f"""Search the internet RIGHT NOW for the 30 most recent CDI job offers in France and Luxembourg 
related to financial security, AML/CFT compliance, LCB-FT analysis, and credit analysis.

Search on: LinkedIn Jobs, France Travail, Indeed France, Welcome to the Jungle, APEC, Jobs.lu

Keywords to search: {', '.join(keywords)}
Location: {location}
Contract type: CDI only
Sort by: Most recent first

For each job found, provide:
1. Job title
2. Company name  
3. Location
4. Direct application URL
5. Publication date (if available)
6. Brief description (1-2 sentences)

Format each job as JSON. Return a JSON array of job objects with keys:
title, company, location, url, date, description

Search NOW and return ONLY real current offers, not examples."""

    try:
        headers = {
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
            "anthropic-beta": "web-search-2025-03-05",
        }
        body = {
            "model": "claude-sonnet-4-20250514",
            "max_tokens": 4000,
            "tools": [{"type": "web_search_20250305", "name": "web_search"}],
            "messages": [{"role": "user", "content": prompt}]
        }
        r = _req.post("https://api.anthropic.com/v1/messages",
                      headers=headers, json=body, timeout=90)
        data = r.json()

        # Extract text from response
        full_text = ""
        for block in data.get("content", []):
            if block.get("type") == "text":
                full_text += block.get("text", "")

        # Try to parse JSON from response
        jobs = []
        if "[" in full_text and "]" in full_text:
            start = full_text.find("[")
            end   = full_text.rfind("]") + 1
            try:
                jobs = json.loads(full_text[start:end])
            except:
                pass

        # If JSON parsing failed, parse structured text
        if not jobs and full_text:
            lines = full_text.split("\n")
            current = {}
            for line in lines:
                line = line.strip()
                if not line: continue
                for key_fr, key_en in [("titre","title"),("entreprise","company"),
                                        ("localisation","location"),("url","url"),
                                        ("date","date"),("description","description")]:
                    for k in [key_fr, key_en, key_fr.capitalize(), key_en.capitalize()]:
                        if line.lower().startswith(k+":") or line.lower().startswith(f'"{k}"'):
                            val = line.split(":",1)[-1].strip().strip('"').strip(",")
                            current[key_en] = val
                            break
                if len(current) >= 4:
                    jobs.append(current); current = {}

        return jobs[:30]

    except Exception as e:
        return [{"title": f"Erreur: {str(e)}", "company": "", "location": "",
                 "url": "", "date": "", "description": ""}]


def generate_lettre_pdf(poste="", entreprise="", secteur="", style_lm="",
                         contexte="", ai_text="", lang="fr") -> bytes:
    from reportlab.pdfgen import canvas as rl_canvas
    buf = io.BytesIO()
    W, H = A4
    d = _get_cv_data(lang)
    target = poste or ("Analyste LCB-FT / Compliance Officer" if lang=="fr" else "AML/CFT Analyst / Compliance Officer")
    ent    = entreprise or ("[Nom de l'établissement]" if lang=="fr" else "[Institution Name]")
    sect   = secteur or ("Finance / Conformité" if lang=="fr" else "Finance / Compliance")
    memoir = d["memoir"]

    ML, MR, MT, MB = 50, 50, 42, 30
    TW = W - ML - MR
    BLK = (0,0,0); DGR=(0.20,0.20,0.20); MGR=(0.40,0.40,0.40); LGR=(0.62,0.62,0.62); VGR=(0.82,0.82,0.82)
    FS = 12; LH = 16.5

    c = rl_canvas.Canvas(buf, pagesize=A4)
    def f(fn,fs): c.setFont(fn,fs)
    def sf(*rgb): c.setFillColorRGB(*rgb)
    def line(y,w=0.6,col=BLK,x1=ML,x2=None):
        if x2 is None: x2=ML+TW
        c.setStrokeColorRGB(*col); c.setLineWidth(w); c.line(x1,y,x2,y)
    def wrap(txt,x,y,mw,fn,fs,lh,col=BLK,ind=0):
        return _wrap_canvas(c,txt,x,y,mw,fn,fs,lh,col,ind)

    y = H - MT

    # EN-TÊTE CENTRÉ
    f("Times-Bold",14); sf(*BLK)
    nw = _sw(d["name"].upper(),"Times-Bold",14)
    c.drawString((W-nw)/2, y, d["name"].upper()); y -= 13
    f("Times-Italic",FS); sf(*DGR)
    tw = _sw(d["title"],"Times-Italic",FS)
    c.drawString((W-tw)/2, y, d["title"]); y -= 11
    f("Times-Roman",9); sf(*MGR)
    row1 = f"{d['loc']}  ·  {d['phone']}  ·  {d['email']}"
    row2 = f"{d['lk']}  ·  Podcast INCLUTECH"
    c.drawString((W-_sw(row1,"Times-Roman",9))/2, y, row1); y -= 10
    c.drawString((W-_sw(row2,"Times-Roman",9))/2, y, row2); y -= 10
    line(y, 0.9, BLK); y -= 2; line(y, 0.3, VGR); y -= 14

    # DESTINATAIRE + DATE
    if lang=="fr":
        f("Times-Roman",FS); sf(*BLK)
        c.drawString(ML, y, "Direction des Ressources Humaines")
        f("Times-Italic",FS); sf(*MGR)
        c.drawRightString(ML+TW, y, "Paris, le 14 mars 2026"); y -= 11
        f("Times-Bold",FS); sf(*BLK)
        c.drawString(ML, y, ent); y -= 11
        f("Times-Roman",FS); sf(*DGR)
        c.drawString(ML, y, sect); y -= 16
        obj_lbl = "Objet : "
        f("Times-Bold",FS); sf(*BLK)
        c.drawString(ML, y, obj_lbl)
        f("Times-Roman",FS)
        c.drawString(ML+_sw(obj_lbl,"Times-Bold",FS), y, f"Candidature — {target}"); y -= 14
    else:
        f("Times-Roman",FS); sf(*BLK)
        c.drawString(ML, y, "Human Resources Department")
        f("Times-Italic",FS); sf(*MGR)
        c.drawRightString(ML+TW, y, "Paris, March 14, 2026"); y -= 11
        f("Times-Bold",FS); sf(*BLK)
        c.drawString(ML, y, ent); y -= 11
        f("Times-Roman",FS); sf(*DGR)
        c.drawString(ML, y, sect); y -= 16
        f("Times-Bold",FS); sf(*BLK)
        obj_lbl = "Re: "
        c.drawString(ML, y, obj_lbl)
        f("Times-Roman",FS)
        c.drawString(ML+_sw(obj_lbl,"Times-Bold",FS), y, f"Application — {target}"); y -= 14

    line(y, 0.6, BLK); y -= 14
    f("Times-Roman",FS); sf(*BLK)
    salut = "Madame, Monsieur," if lang=="fr" else "Dear Hiring Manager,"
    c.drawString(ML, y, salut); y -= 16

    # 3 PARAGRAPHES
    if ai_text and len(ai_text) > 200:
        parts = ai_text.split("\n\n")[:3]
        for p_txt in parts:
            if p_txt.strip() and y > MB+20:
                y = wrap(p_txt.strip(), ML, y, TW, "Times-Roman", FS, LH, DGR, 18); y -= 10
    else:
        if lang == "fr":
            p1 = (
                f"Actuellement Analyste LCB-FT en alternance à la Banque Delubac & Cie sur clientèle "
                f"corporate, CIB, asset management et banque privée, et fort d'une expérience de "
                f"Chargé de Conformité chez HiPay SAS (fintech agréée ACPR — risques pays GAFI, "
                f"KYC/PPE/SOE, pays sensibles, Tracfin, Due Diligence), je me permets de vous "
                f"adresser ma candidature au poste de {target} au sein de {ent}. "
                f"Préparant un MBA 2 Finance de Marché (parcours anglais) à l'ESLSCA Paris, "
                f"certifié AMF et CAMS (Sanctions & Embargos), candidat CFA Level I, "
                f"je combine rigueur réglementaire et maîtrise des marchés financiers."
            )
            p2 = (
                f"Mes missions chez Delubac couvrent le traitement des alertes de niveau 2, "
                f"les Examens Renforcés (ER), l'analyse KYC, l'identification des bénéficiaires "
                f"effectifs, la rédaction des déclarations de soupçon (DS) et l'analyse des états "
                f"financiers. Chez HiPay, j'ai analysé les risques sur pays sensibles et complexes "
                f"(sanctions internationales, jeux en ligne, marketplaces), identifié les PPE/PE/SOE "
                f"et répondu aux réquisitions Tracfin. Mon mémoire de 80 pages — "
                f"« {memoir[:60]}{'…' if len(memoir)>60 else ''} » — "
                f"illustre ma capacité à traiter des problématiques à l'intersection de la "
                f"réglementation ESG et de la performance des portefeuilles."
            )
            p3 = (
                f"Lauréat du 1er Prix Hackathon Fintech Générations 2023 (France FinTech / "
                f"Société Générale, Projet Victoria), Board Member d'Erasmus Expertise et "
                f"créateur du podcast INCLUTECH sur l'inclusion financière, j'apporte une "
                f"vision entrepreneuriale complémentaire à la rigueur analytique attendue. "
                f"Convaincu que la conformité LCB-FT est un avantage concurrentiel, "
                f"je souhaite m'investir pleinement dans votre dispositif de sécurité financière. "
                f"Disponible pour un entretien, je reste à votre disposition. "
                f"Dans l'attente, je vous prie d'agréer, Madame, Monsieur, "
                f"l'expression de mes salutations distinguées."
            )
        else:
            p1 = (
                f"Currently a work-study AML/CFT Analyst at Banque Delubac & Cie (Corporate, CIB, "
                f"Asset Management clients) and with proven Compliance experience at HiPay SAS "
                f"(ACPR-licensed fintech — FATF risk countries, KYC/PEP/SOE, high-risk jurisdictions, "
                f"Tracfin, Due Diligence), I am pleased to apply for the position of {target} "
                f"at {ent}. Currently pursuing an MBA in Market Finance (English track) at "
                f"ESLSCA Paris, AMF and CAMS Certified (Sanctions & Embargoes), and CFA Level I "
                f"Candidate, I combine regulatory expertise with strong financial markets knowledge."
            )
            p2 = (
                f"At Delubac, my responsibilities include Level 2 AML/CFT alert processing, "
                f"Enhanced Due Diligence (EDD), KYC analysis, beneficial owner identification, "
                f"SAR drafting and financial statement analysis for client health assessment. "
                f"At HiPay, I analyzed FATF risk countries including sanctioned jurisdictions "
                f"(online gaming, marketplace clients), screened for PEP/SOE entities, and "
                f"responded to Tracfin and judicial requisitions. My 80-page thesis — "
                f"« {memoir[:60]}{'…' if len(memoir)>60 else ''} » — "
                f"demonstrates my ability to address complex issues at the intersection of "
                f"ESG regulation and portfolio performance."
            )
            p3 = (
                f"As the 1st Prize winner of the Fintech Générations Hackathon 2023 "
                f"(France FinTech / Société Générale, Project Victoria), Board Member of "
                f"Erasmus Expertise, and creator of the INCLUTECH podcast on financial inclusion, "
                f"I bring an entrepreneurial perspective that complements the analytical rigor "
                f"required in compliance roles. I believe AML/CFT compliance is a genuine "
                f"competitive advantage, and I am committed to contributing fully to your "
                f"financial security framework. I would welcome the opportunity to discuss "
                f"my application further. Yours sincerely,"
            )

        for p_txt in [p1, p2, p3]:
            y = wrap(p_txt, ML, y, TW, "Times-Roman", FS, LH, DGR, 18); y -= 10

    # SIGNATURE
    y -= 4
    f("Times-Bold",FS+1); sf(*BLK); c.drawString(ML, y, d["name"]); y -= 12
    f("Times-Italic",FS); sf(*DGR); c.drawString(ML, y, d["title"]); y -= 10
    f("Times-Roman",9); sf(*MGR); c.drawString(ML, y, f"{d['phone']}  ·  {d['email']}"); y -= 9
    sf(*DGR); c.drawString(ML, y, d["lk"])

    # FOOTER PJ
    c.setStrokeColorRGB(*VGR); c.setLineWidth(0.3)
    c.line(ML, MB+8, ML+TW, MB+8)
    f("Times-Italic",8.5); sf(*LGR)
    pj = ("PJ : CV  ·  Mémoire disponible sur demande  ·  Hackathon Victoria  ·  Podcast INCLUTECH"
          if lang=="fr" else
          "Enc.: CV  ·  Thesis available on request  ·  Hackathon Victoria  ·  Podcast INCLUTECH")
    pjw = _sw(pj,"Times-Italic",8.5)
    c.drawString((W-pjw)/2, MB-2, pj)

    c.save(); buf.seek(0); return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# LETTRE DE MOTIVATION DOCX
# ═══════════════════════════════════════════════════════════════════════════════
def generate_lettre_docx(poste="", entreprise="", secteur="", lang="fr") -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Generate PDF text first, then convert to docx
    # For simplicity, use the same paragraph content
    d = _get_cv_data(lang)
    target = poste or ("Analyste LCB-FT" if lang=="fr" else "AML/CFT Analyst")
    ent = entreprise or ("[Établissement]" if lang=="fr" else "[Institution]")
    memoir = d["memoir"]

    doc = Document()
    for sec_d in doc.sections:
        sec_d.top_margin    = Inches(0.65)
        sec_d.bottom_margin = Inches(0.55)
        sec_d.left_margin   = Inches(0.75)
        sec_d.right_margin  = Inches(0.75)

    def p(text, bold=False, italic=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT,
           color=None, sb=0, sa=4):
        pg = doc.add_paragraph()
        pg.alignment = align
        pg.paragraph_format.space_before = Pt(sb)
        pg.paragraph_format.space_after  = Pt(sa)
        r = pg.add_run(text)
        r.bold=bold; r.italic=italic; r.font.size=Pt(size)
        if color: r.font.color.rgb=RGBColor(*color)
        return pg

    p(d["name"].upper(), bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    p(d["title"], italic=True, size=10, color=(50,50,50), align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
    p(f"{d['loc']}  ·  {d['phone']}  ·  {d['email']}", size=9, color=(80,80,80),
      align=WD_ALIGN_PARAGRAPH.CENTER, sa=1)
    p(f"{d['lk']}  ·  Podcast INCLUTECH", size=9, color=(80,80,80),
      align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)

    if lang=="fr":
        p("Direction des Ressources Humaines", size=10, sa=1)
        p(ent, bold=True, size=10, sa=1)
        p("Paris, le 14 mars 2026", italic=True, size=9.5, color=(80,80,80), sa=10)
        p(f"Objet : Candidature — {target}", bold=True, size=10, sa=10)
        p("Madame, Monsieur,", size=10, sa=8)
        paras = [
            (f"Actuellement Analyste LCB-FT en alternance à la Banque Delubac & Cie sur clientèle "
             f"corporate, CIB, asset management, et fort d'une expérience de Chargé de Conformité "
             f"chez HiPay SAS (fintech agréée ACPR — risques pays GAFI, KYC/PPE/SOE, pays sensibles, "
             f"Tracfin), je me permets de vous adresser ma candidature au poste de {target} au "
             f"sein de {ent}. Préparant un MBA 2 Finance de Marché à l'ESLSCA Paris, certifié AMF "
             f"et CAMS (Sanctions & Embargos), candidat CFA Level I, je combine rigueur "
             f"réglementaire et maîtrise des marchés financiers."),
            (f"Mes missions couvrent le traitement des alertes de niveau 2 LCB-FT, les Examens "
             f"Renforcés (ER), l'analyse KYC, la rédaction des déclarations de soupçon (DS) et "
             f"l'analyse des états financiers clients. Mon mémoire de 80 pages — "
             f"« {memoir[:70]}{'…' if len(memoir)>70 else ''} » — illustre ma capacité à traiter "
             f"des problématiques à l'intersection de la réglementation ESG et de la performance "
             f"des portefeuilles."),
            (f"Lauréat du 1er Prix Hackathon Fintech Générations 2023 (France FinTech / Société "
             f"Générale), Board Member d'Erasmus Expertise et créateur du podcast INCLUTECH, "
             f"j'apporte une vision entrepreneuriale complémentaire. Disponible pour un entretien, "
             f"je vous prie d'agréer l'expression de mes salutations distinguées."),
        ]
    else:
        p("Human Resources Department", size=10, sa=1)
        p(ent, bold=True, size=10, sa=1)
        p("Paris, March 14, 2026", italic=True, size=9.5, color=(80,80,80), sa=10)
        p(f"Re: Application — {target}", bold=True, size=10, sa=10)
        p("Dear Hiring Manager,", size=10, sa=8)
        paras = [
            (f"Currently a work-study AML/CFT Analyst at Banque Delubac & Cie (Corporate, CIB, "
             f"Asset Management) and with Compliance experience at HiPay SAS (ACPR-licensed fintech — "
             f"FATF risk countries, KYC/PEP/SOE, high-risk jurisdictions, Tracfin), I am applying "
             f"for the position of {target} at {ent}. Pursuing an MBA in Market Finance at ESLSCA "
             f"Paris, AMF and CAMS Certified, CFA Level I Candidate."),
            (f"At Delubac, I process Level 2 AML alerts, conduct Enhanced Due Diligence, perform KYC "
             f"analysis, draft Suspicious Activity Reports and analyze client financial statements. "
             f"My thesis — « {memoir[:70]}{'…' if len(memoir)>70 else ''} » — addresses ESG "
             f"regulation and portfolio performance."),
            (f"As 1st Prize winner of Fintech Générations Hackathon 2023, Board Member at Erasmus "
             f"Expertise and creator of the INCLUTECH podcast, I bring entrepreneurial perspective "
             f"to complement analytical rigor. I welcome the opportunity to discuss my application. "
             f"Yours sincerely,"),
        ]

    for para_txt in paras:
        p(para_txt, size=10, color=(30,30,30), sa=8)

    p(d["name"], bold=True, size=11, sa=2)
    p(d["title"], italic=True, size=10, color=(50,50,50), sa=10)
    p("PJ : CV  ·  Mémoire disponible sur demande" if lang=="fr" else
      "Enc.: CV  ·  Thesis available on request",
      italic=True, size=9, color=(100,100,100), sa=0)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ═══════════════════════════════════════════════════════════════════════════════
# JOB SCRAPER — CDI Sécurité Financière (France + Luxembourg)
# ═══════════════════════════════════════════════════════════════════════════════
def _scrape_jobs():
    """Return structured job board queries for AML/compliance CDI positions."""
    boards = [
        {
            "source": "LinkedIn Jobs",
            "icon": "💼",
            "queries": [
                {"title":"Analyste LCB-FT / AML — CDI Île-de-France",
                 "url":"https://www.linkedin.com/jobs/search/?keywords=analyste+LCB-FT+AML&location=%C3%8Ele-de-France&f_JT=F&sortBy=DD&f_TPR=r604800",
                 "keywords":["LCB-FT","AML","CDI"]},
                {"title":"Compliance Officer — CDI Paris",
                 "url":"https://www.linkedin.com/jobs/search/?keywords=compliance+officer+securite+financiere&location=Paris&f_JT=F&sortBy=DD&f_TPR=r604800",
                 "keywords":["Compliance","LCB-FT","CDI"]},
                {"title":"AML/KYC Analyst — CDI Luxembourg",
                 "url":"https://www.linkedin.com/jobs/search/?keywords=AML+KYC+analyst+compliance&location=Luxembourg&f_JT=F&sortBy=DD",
                 "keywords":["AML","KYC","Luxembourg"]},
                {"title":"Chargé de Conformité — CDI Fintech",
                 "url":"https://www.linkedin.com/jobs/search/?keywords=charge+conformite+fintech+LCB-FT&location=France&f_JT=F&sortBy=DD",
                 "keywords":["Conformité","Fintech","CDI"]},
                {"title":"Analyste Crédit / Risques — CDI",
                 "url":"https://www.linkedin.com/jobs/search/?keywords=analyste+credit+risques+financier&location=France&f_JT=F&sortBy=DD",
                 "keywords":["Crédit","Risques","CDI"]},
            ]
        },
        {
            "source": "France Travail (Pôle Emploi)",
            "icon": "🇫🇷",
            "queries": [
                {"title":"Analyste LCB-FT / Sécurité Financière — CDI",
                 "url":"https://candidat.francetravail.fr/offres/recherche?motsCles=analyste+LCB-FT+s%C3%A9curit%C3%A9+financi%C3%A8re&typeContrat=CDI&lieux=75&sort=1",
                 "keywords":["LCB-FT","AML","CDI"]},
                {"title":"Compliance Officer AML — CDI",
                 "url":"https://candidat.francetravail.fr/offres/recherche?motsCles=compliance+officer+AML&typeContrat=CDI&lieux=75&sort=1",
                 "keywords":["Compliance","AML","CDI"]},
                {"title":"Analyste Conformité Bancaire — CDI",
                 "url":"https://candidat.francetravail.fr/offres/recherche?motsCles=analyste+conformit%C3%A9+bancaire&typeContrat=CDI&sort=1",
                 "keywords":["Conformité","Bancaire","CDI"]},
                {"title":"Analyste Crédit — CDI Paris",
                 "url":"https://candidat.francetravail.fr/offres/recherche?motsCles=analyste+cr%C3%A9dit&typeContrat=CDI&lieux=75&sort=1",
                 "keywords":["Crédit","Analyse","CDI"]},
                {"title":"KYC / Due Diligence Analyst — CDI",
                 "url":"https://candidat.francetravail.fr/offres/recherche?motsCles=KYC+due+diligence&typeContrat=CDI&sort=1",
                 "keywords":["KYC","Due Diligence","CDI"]},
            ]
        },
        {
            "source": "Welcome to the Jungle",
            "icon": "🌿",
            "queries": [
                {"title":"Compliance / AML LCB-FT — CDI",
                 "url":"https://www.welcometothejungle.com/fr/jobs?query=compliance+AML+LCB-FT&refinementList%5Bcontract_type_names%5D%5B%5D=CDI&sortBy=publishedAt_desc",
                 "keywords":["AML","LCB-FT","CDI"]},
                {"title":"Analyste Sécurité Financière — CDI",
                 "url":"https://www.welcometothejungle.com/fr/jobs?query=analyste+s%C3%A9curit%C3%A9+financi%C3%A8re&refinementList%5Bcontract_type_names%5D%5B%5D=CDI&sortBy=publishedAt_desc",
                 "keywords":["Sécurité Financière","CDI"]},
                {"title":"Chargé de Conformité Fintech — CDI",
                 "url":"https://www.welcometothejungle.com/fr/jobs?query=charg%C3%A9+conformit%C3%A9+fintech&refinementList%5Bcontract_type_names%5D%5B%5D=CDI&sortBy=publishedAt_desc",
                 "keywords":["Conformité","Fintech","CDI"]},
                {"title":"Compliance Officer Banque — CDI",
                 "url":"https://www.welcometothejungle.com/fr/jobs?query=compliance+officer+banque&refinementList%5Bcontract_type_names%5D%5B%5D=CDI&sortBy=publishedAt_desc",
                 "keywords":["Compliance","Banque","CDI"]},
            ]
        },
        {
            "source": "APEC (Cadres)",
            "icon": "📋",
            "queries": [
                {"title":"Conformité AML/LCB-FT — CDI cadre",
                 "url":"https://www.apec.fr/candidat/recherche-emploi.html/emploi?motsCles=conformit%C3%A9+AML+LCB-FT&typeContrat=CDI&sortsBase=DATE_PUBLICATION&sortsOrder=DECREASE",
                 "keywords":["AML","LCB-FT","CDI"]},
                {"title":"Analyste Risques Financiers — CDI",
                 "url":"https://www.apec.fr/candidat/recherche-emploi.html/emploi?motsCles=analyste+risques+financiers&typeContrat=CDI&sortsBase=DATE_PUBLICATION&sortsOrder=DECREASE",
                 "keywords":["Risques","Analyse","CDI"]},
                {"title":"Analyste Crédit Senior — CDI",
                 "url":"https://www.apec.fr/candidat/recherche-emploi.html/emploi?motsCles=analyste+cr%C3%A9dit+senior&typeContrat=CDI&sortsBase=DATE_PUBLICATION&sortsOrder=DECREASE",
                 "keywords":["Crédit","Senior","CDI"]},
            ]
        },
        {
            "source": "Indeed France",
            "icon": "🔍",
            "queries": [
                {"title":"Analyste LCB-FT Paris — CDI",
                 "url":"https://fr.indeed.com/jobs?q=analyste+LCB-FT+s%C3%A9curit%C3%A9+financi%C3%A8re&l=Paris&jt=fulltime&sort=date",
                 "keywords":["LCB-FT","CDI"]},
                {"title":"AML Compliance Officer — CDI",
                 "url":"https://fr.indeed.com/jobs?q=AML+compliance+officer&l=Paris&jt=fulltime&sort=date",
                 "keywords":["AML","Compliance","CDI"]},
                {"title":"KYC Analyst — CDI Île-de-France",
                 "url":"https://fr.indeed.com/jobs?q=KYC+analyst+conformit%C3%A9&l=%C3%8Ele-de-France&jt=fulltime&sort=date",
                 "keywords":["KYC","Analyste","CDI"]},
                {"title":"Analyste Crédit — CDI Paris",
                 "url":"https://fr.indeed.com/jobs?q=analyste+cr%C3%A9dit+risques&l=Paris&jt=fulltime&sort=date",
                 "keywords":["Crédit","Analyse","CDI"]},
            ]
        },
        {
            "source": "Luxembourg Jobs",
            "icon": "🇱🇺",
            "queries": [
                {"title":"Compliance / AML Officer — Luxembourg CDI",
                 "url":"https://www.jobs.lu/en/job-search/?search=compliance+AML+officer&sort=date&contract=permanent",
                 "keywords":["AML","Compliance","Luxembourg"]},
                {"title":"KYC / AML Analyst — Luxembourg",
                 "url":"https://www.jobs.lu/en/job-search/?search=KYC+AML+analyst&sort=date",
                 "keywords":["KYC","AML","Luxembourg"]},
                {"title":"Financial Crime Analyst — Luxembourg",
                 "url":"https://www.jobs.lu/en/job-search/?search=financial+crime+analyst&sort=date",
                 "keywords":["Financial Crime","AML","Luxembourg"]},
                {"title":"Compliance Risk — Luxembourg CDI",
                 "url":"https://www.jobs.lu/en/job-search/?search=compliance+risk+financial&sort=date",
                 "keywords":["Compliance","Risk","Luxembourg"]},
            ]
        },
        {
            "source": "Moovijob (Luxembourg)",
            "icon": "🏛️",
            "queries": [
                {"title":"AML / Conformité — Luxembourg CDI",
                 "url":"https://www.moovijob.com/emploi?keywords=AML+conformite&country=lu&contract=cdi",
                 "keywords":["AML","Luxembourg","CDI"]},
                {"title":"Analyste Risques / Compliance — Luxembourg",
                 "url":"https://www.moovijob.com/emploi?keywords=analyste+risques+compliance&country=lu",
                 "keywords":["Risques","Compliance","Luxembourg"]},
            ]
        },
    ]
    return boards

def generate_bio_pdf() -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=1.7*cm, rightMargin=1.7*cm,
        topMargin=2.0*cm, bottomMargin=1.4*cm)

    story = []

    # ── PAGE 1 : COVER + PHASES 1–3 ──────────────────────────────────
    # Title block
    t_data = [[
        Paragraph(
            '<font name="Helvetica-Bold" size="26" color="#0D1B2A">CINÉUS</font>'
            '  <font name="Helvetica" size="26" color="#0D1B2A">Mackenson</font>',
            _s('bn', fontSize=26, leading=30, spaceAfter=0)),
        Paragraph(
            '<font color="#C9A84C" size="8">Finance · Compliance · Fintech · Paris</font>',
            _s('bs', fontSize=8, alignment=TA_RIGHT, spaceAfter=0,
               textColor=rl_colors.HexColor('#C9A84C')))
    ]]
    story.append(Table(t_data, colWidths=[11.5*cm, 5.3*cm]))
    story.append(_hr(NAVY_C, 2.0, before=3, after=4))

    story.append(Paragraph(
        '<i>"De Port-au-Prince à Levallois-Perret — Compliance Officer en fintech de paiement, '
        'analyste LCB-FT en banque privée, lauréat d\'un hackathon France FinTech, créateur d\'un '
        'podcast sur l\'inclusion financière. Un parcours construit avec intention."</i>',
        _s('q', fontSize=8.5, fontName='Helvetica-Oblique', textColor=NAVY_M_C,
           leading=12.5, leftIndent=6, spaceAfter=8, backColor=CREAM_C)))

    # ── KPI ROW ──────────────────────────────────────────────────────
    kpis = [
        ("💳", "HiPay", "Compliance Officer\n2024–2025"),
        ("🏦", "Delubac", "Analyste LCB-FT\n2023–2024"),
        ("🥇", "Hackathon 2023", "1er Prix · Victoria\nFrance FinTech"),
        ("🎙", "INCLUTECH", "Podcast Fintech\nSpotify & Apple"),
        ("🌍", "Erasmus", "Board Member\nONG internationale"),
    ]
    kpi_cells = []
    for icon, ttl, sub in kpis:
        kpi_cells.append(Paragraph(
            f'<font size="12">{icon}</font><br/>'
            f'<b><font size="7.5" color="#C9A84C">{ttl}</font></b><br/>'
            f'<font size="6.5" color="#888888">{sub.replace(chr(10),"<br/>")}</font>',
            _s('kc', fontSize=6.5, leading=9.5, alignment=TA_CENTER)))
    kpi_t = Table([kpi_cells], colWidths=[3.33*cm]*5)
    kpi_t.setStyle(TableStyle([
        ('ALIGN',(0,0),(-1,-1),'CENTER'), ('VALIGN',(0,0),(-1,-1),'TOP'),
        ('TOPPADDING',(0,0),(-1,-1),5), ('BOTTOMPADDING',(0,0),(-1,-1),5),
        ('ROWBACKGROUNDS',(0,0),(-1,-1),[CREAM_C]),
        ('LINEAFTER',(0,0),(3,-1),0.3,GRAY_L_C),
    ]))
    story.append(kpi_t)
    story.append(Spacer(1, 6))

    def phase_block(label, title, color_hex, paras_html, sources):
        """Render one phase block with colored left bar."""
        # Sources line
        src_parts = []
        for name_s, url_s in sources:
            src_parts.append(
                f'<font color="#4A9EFF"><u>{name_s}</u></font> '
                f'<font size="6" color="#888888">({url_s})</font>')
        src_line = "  ·  ".join(src_parts)

        body_paras = []
        for txt in paras_html:
            body_paras.append(
                Paragraph(txt, _s('pb', fontSize=8.2, leading=12.5,
                                   alignment=TA_JUSTIFY, spaceAfter=4)))
        if src_line:
            body_paras.append(Paragraph(
                f'<font size="6.5" color="#888888"><i>Sources : {src_line}</i></font>',
                _s('src', fontSize=6.5, spaceAfter=0)))

        header = Table([[
            "",
            Paragraph(
                f'<font color="{color_hex}" size="6.5"><b>{label.upper()}</b></font><br/>'
                f'<font size="10" color="#0D1B2A"><b>{title}</b></font>',
                _s('ph', fontSize=10, leading=13.5, spaceAfter=0))
        ]], colWidths=[0.22*cm, 16.3*cm])
        header.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(0,-1), rl_colors.HexColor(color_hex)),
            ('BACKGROUND',(1,0),(1,-1), CREAM_C),
            ('TOPPADDING',(0,0),(-1,-1),5), ('BOTTOMPADDING',(0,0),(-1,-1),5),
            ('LEFTPADDING',(1,0),(1,-1),7), ('RIGHTPADDING',(1,0),(1,-1),5),
        ]))

        inner_rows = [[p] for p in body_paras]
        inner = Table(inner_rows, colWidths=[16.6*cm])
        inner.setStyle(TableStyle([
            ('TOPPADDING',(0,0),(-1,-1),0), ('BOTTOMPADDING',(0,0),(-1,-1),0),
            ('LEFTPADDING',(0,0),(-1,-1),8), ('RIGHTPADDING',(0,0),(-1,-1),0),
        ]))

        return KeepTogether([header, inner, Spacer(1, 5)])

    # ── PHASE 1 ──────────────────────────────────────────────────────
    story.append(phase_block(
        "Phase 1 · 2014–2019 · Haïti",
        "Racines Académiques & Éveil Entrepreneurial",
        "#4A9EFF",
        [
            "Mackenson Cineus intègre l'<b>INAGHEI</b> (Institut National d'Administration de "
            "Gestion et des Hautes Études Internationales) et la <b>FDSE</b> à l'Université "
            "d'État d'Haïti. Il enseigne à mi-temps en lycée et s'engage au <b>Rotaract Club</b>.",

            "En 2018, il co-fonde <b>ANGAJMAN</b> (projets civiques dans 10 départements) et "
            "représente Haïti au <b>Hult Prize</b> avec le projet <b>WEISS</b> — biogaz et "
            "fertilisant organique à partir de déchets. Il milite au <b>Parlement Haïtien de "
            "la Jeunesse pour l'Eau (PHJEA)</b> pour l'accès équitable à l'eau potable.",
        ],
        [("INAGHEI / UEH", "ueh.edu.ht"), ("Hult Prize", "hultprize.org"),
         ("Rotary / Rotaract", "rotary.org")]))

    # ── PHASE 2 ──────────────────────────────────────────────────────
    story.append(phase_block(
        "Phase 2 · 2019–2021 · Le Mans",
        "Transition Internationale — Finance Européenne & Podcast INCLUTECH",
        "#C9A84C",
        [
            "Mackenson décroche une <b>Licence Banque, Finance & Assurance</b> à l'<b>Université "
            "du Mans</b> — marchés bancaires, risk management, réglementation européenne.",

            "Il rejoint <b>Erasmus Expertise</b> (projet ASPEC) et lance le podcast "
            "<b>INCLUTECH</b> (Spotify & Apple Podcasts) : bi-mensuel, dédié au rôle des "
            "fintechs dans l'inclusion financière des populations non bancarisées.",
        ],
        [("Université du Mans", "univ-lemans.fr"),
         ("Erasmus Expertise", "erasmus-expertise.eu"),
         ("INCLUTECH Spotify", "open.spotify.com/show/5XvFdWYwhHWY3EguIhhf69")]))

    # ── PHASE 3 ──────────────────────────────────────────────────────
    story.append(phase_block(
        "Phase 3 · 2021–2023 · Paris",
        "MBA ESLSCA & 1er Prix Hackathon Fintech Générations 2023",
        "#E67E22",
        [
            "<b>MBA Trading & Marchés Financiers</b> à l'<b>ESLSCA Business School Paris</b> "
            "— gestion de portefeuille, produits dérivés, analyse quantitative, conformité.",

            "Octobre 2023 : <b>1er Prix Hackathon Fintech Générations</b> organisé par "
            "<b>France FinTech, Société Générale, Treezor et Paris&Co</b> avec le projet "
            "<b>Victoria</b> (financement rénovations DPE). Invitation à pitcher à "
            "<b>Fintech R:Evolution</b> (1 500 participants). Nommé <b>Board Member "
            "d'Erasmus Expertise</b>.",
        ],
        [("ESLSCA Paris", "eslsca.fr"), ("France FinTech", "francefintech.org"),
         ("Société Générale", "societegenerale.com"), ("Treezor", "treezor.com")]))

    # ── PAGE BREAK ───────────────────────────────────────────────────
    from reportlab.platypus import PageBreak
    story.append(PageBreak())

    # ── PAGE 2 : PHASES 4–5 + SYNTHÈSE ──────────────────────────────
    story.append(phase_block(
        "Phase 4 · 2023–2024 · Paris",
        "Banque Delubac & Cie — Analyste Sécurité Financière LCB-FT",
        "#9B59B6",
        [
            "<b>Banque Delubac & Cie</b> (fondée 1924, banque privée & gestion d'actifs) — "
            "<b>Analyste Sécurité Financière LCB-FT</b> : analyse transactions suspectes "
            "(SAR/STR), évaluation profils de risque KYC/KYB, application directives "
            "<b>AMLD5/AMLD6</b> sous supervision ACPR/GAFI.",

            "Il consolide sa maîtrise des processus LCB-FT en environnement bancaire réglementé "
            "et accompagne des profils MBA ESLSCA dans leur insertion professionnelle.",
        ],
        [("Banque Delubac", "delubac.com"),
         ("ACPR", "acpr.banque-france.fr"), ("GAFI", "fatf-gafi.org")]))

    story.append(phase_block(
        "Phase 5 · 2024–2025 · Levallois-Perret",
        "HiPay — Compliance Officer en Fintech de Paiement",
        "#00B4D8",
        [
            "<b>HiPay</b> est un prestataire de services de paiement omnicanal coté sur "
            "<b>Euronext Growth (ALHYP)</b>, agréé par l'<b>ACPR</b> comme établissement de "
            "paiement. Clients : Veepee, Nocibé, Pizza Hut, Metro.",

            "Mackenson y est <b>Compliance Officer</b> : gestion des alertes niveau 2, "
            "traitement dossiers <b>KYS</b> (Know Your Structure), préparation des "
            "<b>déclarations Tracfin</b> (CRF), vérification questionnaires partenaires, "
            "<b>veille réglementaire</b> (CMF, ACPR, GAFI) avec plans d'action associés.",

            "Cette expérience en fintech de paiement complète idéalement son expertise "
            "bancaire : Mackenson maîtrise désormais la conformité dans les établissements "
            "bancaires traditionnels ET les acteurs tech du paiement numérique.",
        ],
        [("HiPay", "hipay.com"), ("ACPR", "acpr.banque-france.fr"),
         ("Tracfin", "economie.gouv.fr/tracfin"), ("GAFI", "fatf-gafi.org")]))

    # ── SYNTHÈSE FINALE ───────────────────────────────────────────────
    story.append(_hr(NAVY_C, 1.5, before=4, after=4))
    story.append(Paragraph(
        '<b><font color="#0D1B2A" size="10">Synthèse — Présences en ligne & Sources</font></b>',
        _s('sh2', fontSize=10, spaceAfter=6)))

    links = [
        ("LinkedIn",    "linkedin.com/in/mackenson-cineus",
         "500+ connexions · Réseau finance/fintech franco-international"),
        ("HiPay",       "hipay.com",
         "Employeur 2024–2025 · Paiement omnicanal · Euronext Growth ALHYP"),
        ("Banque Delubac","delubac.com",
         "Employeur 2023–2024 · Banque privée fondée 1924"),
        ("INCLUTECH Spotify","open.spotify.com/show/5XvFdWYwhHWY3EguIhhf69",
         "Podcast bi-mensuel · Finance & Inclusion · Spotify & Apple Podcasts"),
        ("France FinTech","francefintech.org",
         "Lauréat Hackathon 2023 · Membre actif · Fintech R:Evolution"),
        ("ESLSCA Paris","eslsca.fr","MBA Trading & Marchés Financiers 2021–2023"),
        ("Université du Mans","univ-lemans.fr","Licence BFA 2019–2021"),
        ("Erasmus Expertise","erasmus-expertise.eu","Board Member · ONG internationale"),
        ("Instagram","instagram.com/mackenson_cineus","@mackenson_cineus"),
        ("ACPR","acpr.banque-france.fr","Régulateur bancaire/paiement · Référence LCB-FT"),
        ("Tracfin","economie.gouv.fr/tracfin","Cellule de Renseignement Financier française"),
        ("GAFI / FATF","fatf-gafi.org","Normes mondiales anti-blanchiment"),
    ]
    link_rows = []
    for i in range(0, len(links), 2):
        row = []
        for j in range(2):
            if i+j < len(links):
                n, u, d = links[i+j]
                row.append(Paragraph(
                    f"<b>{n}</b>  <font color='#4A9EFF' size='7'>{u}</font><br/>"
                    f"<font size='6.8' color='#666666'>{d}</font>",
                    _s('lnk', fontSize=7, leading=10, spaceAfter=0)))
            else:
                row.append(Paragraph("", _s('x')))
        link_rows.append(row)

    link_t = Table(link_rows, colWidths=[8.5*cm, 8.5*cm])
    link_t.setStyle(TableStyle([
        ('TOPPADDING',(0,0),(-1,-1),4), ('BOTTOMPADDING',(0,0),(-1,-1),4),
        ('VALIGN',(0,0),(-1,-1),'TOP'),
        ('ROWBACKGROUNDS',(0,0),(-1,-1),[CREAM_C, rl_colors.white]),
        ('LINEAFTER',(0,0),(0,-1),0.3, GRAY_L_C),
        ('INNERGRID',(0,0),(-1,-1),0.2, rl_colors.HexColor('#EEEEEE')),
        ('LEFTPADDING',(0,0),(-1,-1),6), ('RIGHTPADDING',(0,0),(-1,-1),6),
    ]))
    story.append(link_t)

    doc.build(story, onFirstPage=_page_frame, onLaterPages=_page_frame)
    buf.seek(0); return buf.read()


# ─── JOB SCRAPER HELPERS ───────────────────────────────────────────────────

# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR — 12 PAGES
# ═══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div style='text-align:center;padding:16px 0 8px;'>
      <div style='width:72px;height:72px;border-radius:50%;
                  background:linear-gradient(135deg,#C9A84C,#4A9EFF);
                  display:flex;align-items:center;justify-content:center;
                  font-size:1.8rem;margin:0 auto 10px;
                  border:2.5px solid rgba(201,168,76,0.5);'>MC</div>
      <div style='font-family:Playfair Display,serif;font-size:1rem;font-weight:700;color:#C9A84C;'>
        Mackenson CINÉUS</div>
      <div style='font-size:0.75rem;color:#F5F0E8;opacity:0.65;margin-top:3px;'>
        Analyste LCB-FT · Compliance · Finance</div>
    </div>
    """, unsafe_allow_html=True)
    st.divider()

    page = st.radio("Navigation", [
        "🏠  Profil & Identité",
        "📖  Biographie Narrative",
        "📄  Biographie Courte",
        "⏱   Parcours & Compétences",
        "📊  Visualisations",
        "🌍  Carte Géographique",
        "🔗  Plateformes & Liens",
        "📥  Générateur de Documents",
        "🔎  Offres CDI (Temps Réel)",
        "🤖  Assistant IA Carrière",
        "🔧  Outils Compliance",
        "✏️   Édition du Profil",
    ], label_visibility="collapsed")

    st.divider()
    st.markdown("<div style='font-size:0.78rem;color:#C9A84C;font-weight:600;margin-bottom:5px;'>🔑 Clé API Anthropic</div>", unsafe_allow_html=True)
    st.markdown("<div style='font-size:0.7rem;opacity:0.55;margin-bottom:6px;'>Pour l'assistant IA et la génération LM</div>", unsafe_allow_html=True)
    api_key_input = st.text_input("", type="password", placeholder="sk-ant-...",
                                   label_visibility="collapsed", key="anthropic_api_key")
    if api_key_input:
        st.markdown("<div style='font-size:0.7rem;color:#2ECC71;'>✅ Clé active</div>", unsafe_allow_html=True)
    else:
        st.markdown("<div style='font-size:0.7rem;color:#E67E22;'>⚠️ Non renseignée</div>", unsafe_allow_html=True)
    st.divider()
    st.markdown("<div style='font-size:0.72rem;opacity:0.45;text-align:center;padding:6px;'>© 2025 Mackenson CINÉUS</div>", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — PROFIL & IDENTITÉ
# ═══════════════════════════════════════════════════════════════════════════════
if page == "🏠  Profil & Identité":
    P = _P()
    st.markdown(_EXTRA_CSS, unsafe_allow_html=True)

    st.markdown(f"""
    <div class="hero-banner">
      <div style="display:flex;align-items:center;gap:20px;flex-wrap:wrap;">
        <div style="width:68px;height:68px;border-radius:50%;
                    background:linear-gradient(135deg,#C9A84C,#4A9EFF);
                    display:flex;align-items:center;justify-content:center;
                    font-size:1.7rem;border:2.5px solid rgba(201,168,76,0.5);flex-shrink:0;">MC</div>
        <div>
          <p class="hero-name">{P['name']}</p>
          <p class="hero-title">{P['title']} · {P['location']}</p>
        </div>
      </div>
      <div style="margin-top:14px;">
        <span class="hero-badge">🏦 LCB-FT N2 · ER · DS/SAR</span>
        <span class="hero-badge">💳 HiPay · PPE/SOE · GAFI</span>
        <span class="hero-badge">📈 MBA Finance · ESLSCA</span>
        <span class="hero-badge">🥇 Hackathon FinTech 2023</span>
        <span class="hero-badge">✅ AMF · CAMS · CFA L1</span>
      </div>
      <p style="margin-top:18px;max-width:760px;opacity:0.85;line-height:1.85;font-size:0.93rem;">
        {P['summary']}
      </p>
    </div>
    """, unsafe_allow_html=True)

    cols = st.columns(5)
    kpis = [
        ("2", "Postes compliance", "#C9A84C"),
        ("3", "Certifications", "#4A9EFF"),
        ("2", "Diplômes sup.", "#2ECC71"),
        ("🥇", "Hackathon 2023", "#E67E22"),
        ("80p", "Mémoire MBA", "#9B59B6"),
    ]
    for col, (val, lbl, color) in zip(cols, kpis):
        with col:
            st.markdown(f"""
            <div class="metric-card" style="border-top:3px solid {color};">
              <span class="metric-value" style="color:{color};font-size:1.9rem;">{val}</span>
              <div class="metric-label">{lbl}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    c1, c2 = st.columns([2, 1])
    with c1:
        st.markdown('<p class="section-header">Personnalité & Profil type</p>', unsafe_allow_html=True)
        personality = st.session_state.get("edit_personality","")
        st.markdown(f"""
        <div style="background:var(--navy-mid);border-left:4px solid #C9A84C;
                    padding:16px 20px;border-radius:0 10px 10px 0;margin-bottom:16px;">
          <div style="font-size:1rem;font-style:italic;line-height:1.8;opacity:0.9;">
            « {personality} »
          </div>
        </div>""", unsafe_allow_html=True)

        st.markdown("""
        <div style="display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-top:8px;">
          <div style="background:var(--navy-mid);border:1px solid rgba(74,158,255,0.2);
                      border-radius:10px;padding:14px;">
            <div style="font-size:0.82rem;color:#4A9EFF;font-weight:600;margin-bottom:6px;">Type de profil</div>
            <div style="font-size:0.88rem;opacity:0.9;line-height:1.6;">
              <b>ENTJ — Commandant</b><br>Leadership structuré · Vision analytique · Entrepreneuriat
            </div>
          </div>
          <div style="background:var(--navy-mid);border:1px solid rgba(46,204,113,0.2);
                      border-radius:10px;padding:14px;">
            <div style="font-size:0.82rem;color:#2ECC71;font-weight:600;margin-bottom:6px;">Valeur unique</div>
            <div style="font-size:0.88rem;opacity:0.9;line-height:1.6;">
              Banque + Fintech + MBA + Hackathon + ONG + Podcast — profil compliance junior rare
            </div>
          </div>
        </div>""", unsafe_allow_html=True)

    with c2:
        st.markdown('<p class="section-header" style="font-size:1.1rem;">🔗 Liens Rapides</p>', unsafe_allow_html=True)
        for icon, name_l, url in [
            ("💼","LinkedIn","https://linkedin.com/in/mackenson-cineus"),
            ("💳","HiPay","https://hipay.com"),
            ("🏦","Banque Delubac","https://www.delubac.com"),
            ("🎙","INCLUTECH","https://open.spotify.com/show/5XvFdWYwhHWY3EguIhhf69"),
            ("🚀","France FinTech","https://francefintech.org"),
            ("🌍","Erasmus Expertise","https://erasmus-expertise.eu"),
        ]:
            st.markdown(f"""
            <a href="{url}" target="_blank" style="text-decoration:none;">
              <div style="display:flex;align-items:center;gap:10px;
                          background:var(--navy-mid);border:1px solid rgba(255,255,255,0.07);
                          border-radius:8px;padding:8px 12px;margin-bottom:6px;">
                <span>{icon}</span>
                <span style="font-size:0.82rem;color:var(--gold);">{name_l}</span>
                <span style="margin-left:auto;opacity:0.35;font-size:0.7rem;">↗</span>
              </div>
            </a>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — BIOGRAPHIE NARRATIVE (phases détaillées)
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "📖  Biographie Narrative":
    st.markdown(_EXTRA_CSS, unsafe_allow_html=True)
    st.markdown('<p class="section-header">📖 Biographie Narrative — Parcours en 4 phases</p>', unsafe_allow_html=True)
    st.markdown("<p style='opacity:0.7;margin-bottom:20px;line-height:1.7;'>Parcours détaillé avec hyperliens vers les sources officielles.</p>", unsafe_allow_html=True)
    for period in BIO_PERIODS:
        tags_html = "".join(f'<span class="bio-tag">{t}</span>' for t in period["tags"])
        paras_html = "".join(f'<p style="margin:0 0 12px 0;">{p}</p>' for p in period["text"])
        st.markdown(f"""
        <div class="bio-period {period['phase']}">
          <div class="bio-phase-label">{period['label']}</div>
          <div class="bio-period-title">{period['title']}</div>
          <div class="bio-period-years">📍 {period['years']}</div>
          <div class="bio-period-text">{paras_html}</div>
          <div style="margin-top:12px;">{tags_html}</div>
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — BIOGRAPHIE COURTE (3 paragraphes segmentés)
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "📄  Biographie Courte":
    st.markdown('<p class="section-header">📄 Biographie Courte — 3 Dimensions</p>', unsafe_allow_html=True)
    st.markdown("<p style='opacity:0.7;margin-bottom:20px;'>Format synthétique en 3 blocs thématiques : Formation · Expérience · Engagement.</p>", unsafe_allow_html=True)

    sections = [
        ("🎓", "Formation & Certifications", "phase2", BIO_SHORT["formation"]),
        ("💼", "Expérience Professionnelle", "phase4", BIO_SHORT["experience"]),
        ("🌱", "Engagements Social, Environnemental & Entrepreneurial", "phase1", BIO_SHORT["engagement"]),
    ]
    for icon, title_b, phase, text in sections:
        st.markdown(f"""
        <div class="bio-period {phase}">
          <div class="bio-phase-label">{icon} {title_b}</div>
          <div class="bio-period-text" style="margin-top:10px;line-height:1.85;">{text}</div>
        </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — PARCOURS & COMPÉTENCES
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "⏱   Parcours & Compétences":
    st.markdown('<p class="section-header">⏱ Parcours & Compétences</p>', unsafe_allow_html=True)
    c1, c2 = st.columns([3, 2])
    with c1:
        cats = ["Tous"] + sorted(set(t["cat"] for t in TIMELINE))
        fc = st.selectbox("Filtrer", cats, label_visibility="collapsed")
        items = TIMELINE if fc=="Tous" else [t for t in TIMELINE if t["cat"]==fc]
        for item in items:
            url_part = f'<a href="{item["url"]}" target="_blank" style="color:#4A9EFF;font-size:0.75rem;"> 🔗</a>' if item.get("url") else ""
            st.markdown(f"""
            <div class="timeline-item">
              <div class="timeline-year">{item['icon']} {item['year']} · {item['cat']}</div>
              <div class="timeline-title">{item['title']}{url_part}</div>
              <div class="timeline-desc">{item['desc']}</div>
            </div>""", unsafe_allow_html=True)
    with c2:
        live_skills = st.session_state.get("edit_skills", SKILLS)
        st.markdown('<p class="section-header">Compétences</p>', unsafe_allow_html=True)
        for skill, level in live_skills.items():
            st.markdown(f"""
            <div class="skill-bar-container">
              <div class="skill-name">{skill}<span style="float:right;color:#C9A84C;font-weight:600;">{level}%</span></div>
              <div class="skill-bar"><div class="skill-fill" style="width:{level}%;"></div></div>
            </div>""", unsafe_allow_html=True)
        st.markdown('<br><p class="section-header">Langues</p>', unsafe_allow_html=True)
        for lang_l, level in LANGUAGES.items():
            c_lang = "#C9A84C" if "Français" in lang_l or "Créole" in lang_l else "#4A9EFF"
            st.markdown(f"""
            <div class="skill-bar-container">
              <div class="skill-name">{lang_l}<span style="float:right;color:{c_lang};font-weight:600;">{level}%</span></div>
              <div class="skill-bar"><div class="skill-fill" style="width:{level}%;background:linear-gradient(90deg,{c_lang},{c_lang}AA);"></div></div>
            </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 5 — VISUALISATIONS
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "📊  Visualisations":
    st.markdown('<p class="section-header">📊 Visualisations</p>', unsafe_allow_html=True)
    tab1, tab2, tab3, tab4 = st.tabs(["Radar Compétences","Progression Carrière","Répartition","Langues"])

    with tab1:
        live_skills = st.session_state.get("edit_skills", SKILLS)
        sn = list(live_skills.keys()); sv = list(live_skills.values())
        fig = go.Figure(go.Scatterpolar(r=sv+[sv[0]], theta=sn+[sn[0]], fill='toself',
            fillcolor='rgba(201,168,76,0.15)', line=dict(color='#C9A84C', width=2),
            marker=dict(color='#C9A84C', size=7)))
        fig.update_layout(polar=dict(bgcolor='rgba(26,46,66,0.5)',
            radialaxis=dict(visible=True,range=[0,100],gridcolor='rgba(255,255,255,0.1)',color='#F5F0E8'),
            angularaxis=dict(gridcolor='rgba(255,255,255,0.1)',color='#F5F0E8',tickfont=dict(size=10))),
            paper_bgcolor='rgba(0,0,0,0)',font=dict(family='DM Sans',color='#F5F0E8'),
            title=dict(text='Radar des Compétences',font=dict(color='#C9A84C',size=15)),height=460)
        st.plotly_chart(fig, use_container_width=True)

    with tab2:
        df = pd.DataFrame([
            {"A":"2016","N":8,"E":"ACTIVEH · Président"},
            {"A":"2018","N":25,"E":"Hult Prize · WEISS"},
            {"A":"2022","N":48,"E":"Licence BFA · Le Mans"},
            {"A":"2023","N":68,"E":"Hackathon Fintech 🥇"},
            {"A":"2024","N":80,"E":"MBA ESLSCA · HiPay"},
            {"A":"2025","N":92,"E":"Banque Delubac LCB-FT"},
        ])
        fig2 = go.Figure(go.Scatter(x=df["A"],y=df["N"],mode='lines+markers+text',
            line=dict(color='#C9A84C',width=3),
            marker=dict(size=13,color='#C9A84C',line=dict(color='#1A2E42',width=3)),
            text=df["E"],textposition='top center',textfont=dict(size=8.5,color='#F5F0E8'),
            hovertemplate='<b>%{text}</b><br>Niveau: %{y}%<extra></extra>'))
        fig2.update_layout(title=dict(text='Progression de Carrière',font=dict(color='#C9A84C',size=15)),
            yaxis=dict(range=[0,110],gridcolor='rgba(255,255,255,0.07)',color='#F5F0E8',title='Développement (%)'),
            xaxis=dict(color='#F5F0E8',gridcolor='rgba(255,255,255,0.07)'),
            paper_bgcolor='rgba(0,0,0,0)',plot_bgcolor='rgba(13,27,42,0.6)',
            font=dict(family='DM Sans',color='#F5F0E8'),height=420)
        st.plotly_chart(fig2, use_container_width=True)

    with tab3:
        c1, c2 = st.columns(2)
        with c1:
            axes = {"Finance & Conformité":45,"Entrepreneuriat Social":22,"Innovation FinTech":18,"Engagement Citoyen":15}
            fig3 = go.Figure(go.Pie(labels=list(axes.keys()),values=list(axes.values()),hole=0.5,
                marker=dict(colors=['#C9A84C','#4A9EFF','#E67E22','#2ECC71'],line=dict(color='#0D1B2A',width=2)),
                textfont=dict(color='white',size=10)))
            fig3.update_layout(title=dict(text='Axes du parcours',font=dict(color='#C9A84C',size=14)),
                paper_bgcolor='rgba(0,0,0,0)',font=dict(family='DM Sans',color='#F5F0E8'),
                legend=dict(font=dict(color='#F5F0E8'),bgcolor='rgba(0,0,0,0)'),height=340)
            st.plotly_chart(fig3, use_container_width=True)
        with c2:
            countries = {"Haïti":35,"France":55,"Luxembourg":10}
            fig4 = go.Figure(go.Bar(x=list(countries.keys()),y=list(countries.values()),
                marker=dict(color=['#4A9EFF','#C9A84C','#2ECC71'],line=dict(color='#0D1B2A',width=1)),
                text=[f"{v}%" for v in countries.values()],textposition='outside',textfont=dict(color='#F5F0E8')))
            fig4.update_layout(title=dict(text='Répartition géographique',font=dict(color='#C9A84C',size=14)),
                yaxis=dict(range=[0,70],gridcolor='rgba(255,255,255,0.07)',color='#F5F0E8'),
                xaxis=dict(color='#F5F0E8'),paper_bgcolor='rgba(0,0,0,0)',plot_bgcolor='rgba(13,27,42,0.6)',
                font=dict(family='DM Sans',color='#F5F0E8'),height=340)
            st.plotly_chart(fig4, use_container_width=True)

    with tab4:
        fig5 = go.Figure()
        clrs = {'Français (Natif)':'#C9A84C','Anglais (C1)':'#4A9EFF','Créole haïtien':'#2ECC71'}
        for lang_l, lvl in LANGUAGES.items():
            fig5.add_trace(go.Bar(x=[lvl],y=[lang_l],orientation='h',
                marker_color=clrs.get(lang_l,'#C9A84C'),
                text=[f"{lvl}%"],textposition='inside',textfont=dict(color='white',size=13),
                hovertemplate=f'{lang_l}: {lvl}%<extra></extra>'))
        fig5.update_layout(title=dict(text='Langues',font=dict(color='#C9A84C',size=15)),
            xaxis=dict(range=[0,110],gridcolor='rgba(255,255,255,0.07)',color='#F5F0E8'),
            yaxis=dict(color='#F5F0E8'),paper_bgcolor='rgba(0,0,0,0)',plot_bgcolor='rgba(13,27,42,0.6)',
            font=dict(family='DM Sans',color='#F5F0E8'),showlegend=False,height=270,barmode='overlay')
        st.plotly_chart(fig5, use_container_width=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 6 — CARTE GÉOGRAPHIQUE
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "🌍  Carte Géographique":
    st.markdown('<p class="section-header">🌍 Parcours Géographique</p>', unsafe_allow_html=True)
    fig_map = go.Figure()
    coords = list(zip(GEO_DATA['lat'],GEO_DATA['lon']))
    for i in range(len(coords)-1):
        fig_map.add_trace(go.Scattergeo(lat=[coords[i][0],coords[i+1][0]],lon=[coords[i][1],coords[i+1][1]],
            mode='lines',line=dict(width=2,color='rgba(201,168,76,0.5)'),showlegend=False))
    fig_map.add_trace(go.Scattergeo(lat=GEO_DATA['lat'],lon=GEO_DATA['lon'],mode='markers+text',
        marker=dict(size=GEO_DATA['size'],color=['#4A9EFF','#C9A84C','#C9A84C'],
                    line=dict(color='white',width=2),opacity=0.9),
        text=GEO_DATA['locations'],textposition=['bottom right','bottom left','top right'],
        textfont=dict(size=11,color='white'),
        customdata=list(zip(GEO_DATA['events'],GEO_DATA['years'])),
        hovertemplate='<b>%{text}</b><br>%{customdata[0]}<br>%{customdata[1]}<extra></extra>'))
    fig_map.update_geos(projection_type="natural earth",showcoastlines=True,coastlinecolor='rgba(255,255,255,0.2)',
        showland=True,landcolor='rgba(26,46,66,0.8)',showocean=True,oceancolor='rgba(13,27,42,0.9)',
        showcountries=True,countrycolor='rgba(255,255,255,0.1)',center=dict(lat=30,lon=-20),projection_scale=2.5)
    fig_map.update_layout(paper_bgcolor='rgba(0,0,0,0)',font=dict(family='DM Sans',color='#F5F0E8'),
        height=480,margin=dict(l=0,r=0,t=20,b=0),geo=dict(bgcolor='rgba(0,0,0,0)'))
    st.plotly_chart(fig_map, use_container_width=True)
    cols = st.columns(3)
    for col, (loc, event, year) in zip(cols, zip(GEO_DATA['locations'],GEO_DATA['events'],GEO_DATA['years'])):
        with col:
            st.markdown(f"""<div class="timeline-item" style="border-left-color:#4A9EFF;">
              <div class="timeline-year">📍 {year}</div>
              <div class="timeline-title">{loc}</div>
              <div class="timeline-desc">{event}</div>
            </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 7 — PLATEFORMES & LIENS
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "🔗  Plateformes & Liens":
    st.markdown('<p class="section-header">🔗 Plateformes & Présences en ligne</p>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    for i, p in enumerate(PLATFORMS):
        with (c1 if i%2==0 else c2):
            st.markdown(f"""
            <div class="platform-card">
              <div class="platform-icon">{p['icon']}</div>
              <div style="flex:1;">
                <div class="platform-name">{p['name']}</div>
                <div class="platform-desc">{p['desc']}</div>
                <a href="{p['url']}" target="_blank" class="platform-link">{p['label']}</a>
              </div>
            </div>""", unsafe_allow_html=True)
    st.markdown('<br><p class="section-header">📡 Régulateurs de référence</p>', unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    for i, (icon, name_r, desc, url) in enumerate(REGULATORS):
        with (c1 if i%2==0 else c2):
            st.markdown(f"""
            <div class="platform-card">
              <div class="platform-icon">{icon}</div>
              <div style="flex:1;">
                <div class="platform-name">{name_r}</div>
                <div class="platform-desc">{desc}</div>
                <a href="{url}" target="_blank" class="platform-link">Consulter →</a>
              </div>
            </div>""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 8 — GÉNÉRATEUR DE DOCUMENTS (PDF + DOCX, FR + EN)
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "📥  Générateur de Documents":
    st.markdown('<p class="section-header">📥 Générateur de Documents Professionnels</p>', unsafe_allow_html=True)
    tab_cv, tab_lm, tab_bio = st.tabs(["📋 CV", "✉️ Lettre de Motivation", "📖 Biographie PDF"])

    with tab_cv:
        st.markdown("""<div style='background:var(--navy-mid);border:1px solid rgba(201,168,76,0.2);
                    border-radius:10px;padding:16px;margin-bottom:16px;'>
          <b style='color:#C9A84C;'>CV Finance — 1 page A4 complète</b><br>
          <span style='font-size:0.82rem;opacity:0.75;'>Noir/gris · police 9.5 · en-tête centré · 5 missions/poste · certifications · mémoire</span>
        </div>""", unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1: cv_poste = st.text_input("🎯 Poste visé", placeholder="Analyste LCB-FT…", key="cv_p")
        with c2: cv_lang  = st.selectbox("🌐 Langue", ["Français","English"], key="cv_l")
        with c3: cv_fmt   = st.selectbox("📄 Format", ["PDF","Word (.docx)"], key="cv_f")

        if st.button("⬇️ Générer le CV", use_container_width=True, key="btn_cv"):
            lang_code = "fr" if cv_lang == "Français" else "en"
            with st.spinner("Génération…"):
                if cv_fmt == "PDF":
                    data = generate_cv_pdf(poste=cv_poste, lang=lang_code)
                    mime = "application/pdf"
                    fname = f"Cineus_Mackenson_CV_{lang_code.upper()}.pdf"
                else:
                    data = generate_cv_docx(poste=cv_poste, lang=lang_code)
                    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    fname = f"Cineus_Mackenson_CV_{lang_code.upper()}.docx"
            st.success("✅ CV généré !")
            st.download_button(f"📥 Télécharger — {fname}", data, fname, mime, use_container_width=True, key="dl_cv")

    with tab_lm:
        st.markdown("""<div style='background:var(--navy-mid);border:1px solid rgba(74,158,255,0.2);
                    border-radius:10px;padding:16px;margin-bottom:16px;'>
          <b style='color:#4A9EFF;'>Lettre de Motivation — 3 paragraphes, 1 page pleine</b><br>
          <span style='font-size:0.82rem;opacity:0.75;'>Mots-clés AML/LCB-FT intégrés · mémoire mentionné · FR ou EN</span>
        </div>""", unsafe_allow_html=True)
        c1, c2 = st.columns(2)
        with c1:
            lm_poste = st.text_input("🎯 Poste", placeholder="Analyste LCB-FT…", key="lm_p")
            lm_ent   = st.text_input("🏢 Entreprise", placeholder="BNP Paribas…", key="lm_e")
        with c2:
            lm_sect  = st.selectbox("🏦 Secteur", ["Banque","Fintech","Asset Management","Assurance","Audit"], key="lm_s")
            lm_lang  = st.selectbox("🌐 Langue", ["Français","English"], key="lm_l")
        lm_fmt = st.selectbox("📄 Format", ["PDF","Word (.docx)"], key="lm_f")
        use_ai = st.toggle("🤖 Personnaliser avec Claude AI", False, key="lm_ai")
        lm_ctx = st.text_area("Instructions IA", height=60, placeholder="Ex: insister sur CAMS, pays sensibles…", key="lm_ctx")

        if st.button("✨ Générer la Lettre", use_container_width=True, key="btn_lm"):
            if not lm_poste:
                st.error("⚠️ Poste requis")
            else:
                lang_code = "fr" if lm_lang == "Français" else "en"
                ai_content = ""
                if use_ai and st.session_state.get("anthropic_api_key","").strip():
                    with st.spinner("🤖 Claude rédige…"):
                        ss = st.session_state
                        sys_p = "Expert en lettres de motivation finance/compliance. 3 paragraphes denses, mots-clés AML/LCB-FT."
                        usr_p = (f"Rédige 3 paragraphes (séparés par ligne vide) pour une lettre de motivation. "
                                 f"Profil: Analyste LCB-FT Delubac (alertes N2, ER, KYC, DS/SAR), Compliance HiPay "
                                 f"(GAFI, PPE/SOE, pays sensibles, Tracfin), MBA ESLSCA, AMF+CAMS+CFA L1, "
                                 f"Hackathon Fintech 2023, Mémoire ESG. Poste: {lm_poste}. Entreprise: {lm_ent}. "
                                 f"Langue: {'français' if lang_code=='fr' else 'anglais'}. "
                                 f"Instructions: {lm_ctx or 'standard'}. SANS formule d'appel ni signature.")
                        try:
                            import requests as _req
                            hdrs = {"Content-Type":"application/json",
                                    "x-api-key": ss.get("anthropic_api_key",""),
                                    "anthropic-version":"2023-06-01"}
                            body = {"model":"claude-sonnet-4-20250514","max_tokens":1000,
                                    "system":sys_p,"messages":[{"role":"user","content":usr_p}]}
                            r = _req.post("https://api.anthropic.com/v1/messages",headers=hdrs,json=body,timeout=60)
                            data_r = r.json()
                            ai_content = "".join(b.get("text","") for b in data_r.get("content",[]) if b.get("type")=="text")
                        except Exception as e:
                            st.warning(f"IA indisponible: {e}")

                with st.spinner("Génération…"):
                    if lm_fmt == "PDF":
                        data = generate_lettre_pdf(poste=lm_poste, entreprise=lm_ent, secteur=lm_sect,
                                                    ai_text=ai_content, lang=lang_code)
                        mime = "application/pdf"
                        fname = f"Cineus_Mackenson_LM_{lang_code.upper()}.pdf"
                    else:
                        data = generate_lettre_docx(poste=lm_poste, entreprise=lm_ent,
                                                     secteur=lm_sect, lang=lang_code)
                        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        fname = f"Cineus_Mackenson_LM_{lang_code.upper()}.docx"
                st.success("✅ Lettre générée !")
                st.download_button(f"📥 {fname}", data, fname, mime, use_container_width=True, key="dl_lm")

    with tab_bio:
        st.markdown("**Biographie PDF complète (2 pages) avec hyperliens sources.**")
        if st.button("⬇️ Générer Biographie PDF", use_container_width=True, key="btn_bio"):
            with st.spinner("Génération…"):
                bio_data = generate_bio_pdf()
            st.success("✅")
            st.download_button("📥 Cineus_Mackenson_Biographie.pdf", bio_data,
                               "Cineus_Mackenson_Biographie.pdf", "application/pdf",
                               use_container_width=True, key="dl_bio")


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 9 — OFFRES CDI TEMPS RÉEL
# ═══════════════════════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 9 — PLATEFORME DE RECHERCHE D'EMPLOI (SCRAPING TEMPS RÉEL)
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "🔎  Offres CDI (Temps Réel)":

    st.markdown('<p class="section-header">🔎 Plateforme de Recherche d\'Emploi — Sécurité Financière</p>',
                unsafe_allow_html=True)

    # ── INTRO ────────────────────────────────────────────────────────────────
    st.markdown("""
    <div style='background:var(--navy-mid);border:1px solid rgba(201,168,76,0.25);
                border-radius:12px;padding:16px 20px;margin-bottom:20px;'>
      <div style='font-size:0.95rem;font-weight:600;color:#C9A84C;margin-bottom:8px;'>
        🔴 Recherche en temps réel · CDI uniquement · France & Luxembourg
      </div>
      <div style='font-size:0.83rem;opacity:0.8;line-height:1.6;'>
        Utilisez votre <b>clé API Anthropic</b> (barre latérale) pour lancer une recherche réelle
        sur LinkedIn, France Travail, Indeed, Welcome to the Jungle, APEC, Jobs.lu — scraping live.
        <br>Sans clé API → accès direct aux plateformes via liens pré-configurés.
      </div>
    </div>
    """, unsafe_allow_html=True)

    api_key = st.session_state.get("anthropic_api_key", "").strip()

    # ── SEARCH CONTROLS ──────────────────────────────────────────────────────
    st.markdown("### ⚙️ Paramètres de recherche")

    col1, col2, col3 = st.columns([3, 2, 1])
    with col1:
        custom_kw = st.text_input(
            "🔍 Mots-clés supplémentaires",
            placeholder="ex: CAMS, Actimize, Trade Finance, TRACFIN...",
            key="job_custom_kw"
        )
    with col2:
        job_location = st.selectbox(
            "📍 Localisation",
            ["France", "Paris / Île-de-France", "Luxembourg", "France + Luxembourg",
             "Île-de-France", "Lyon", "Remote / Télétravail"],
            key="job_location"
        )
    with col3:
        st.markdown("<div style='height:28px'></div>", unsafe_allow_html=True)
        launch_search = st.button("🚀 Lancer la recherche", use_container_width=True, key="launch_search")

    # Default keyword categories
    st.markdown("### 🏷️ Catégories de recherche")

    kw_categories = {
        "🏦 LCB-FT / AML": ["analyste LCB-FT", "compliance AML", "lutte blanchiment", 
                              "sécurité financière CDI", "AML analyst", "financial crime"],
        "🔍 KYC / Due Diligence": ["KYC analyst CDI", "due diligence compliance",
                                    "chargé conformité KYC", "KYB analyst"],
        "⚖️ Conformité Réglementaire": ["compliance officer CDI", "chargé conformité bancaire",
                                          "conformité réglementaire", "regulatory compliance"],
        "💳 Fintech / PSP": ["compliance fintech CDI", "chargé conformité PSP",
                              "fintech AML", "établissement paiement conformité"],
        "📊 Analyse Crédit": ["analyste crédit CDI Paris", "credit analyst compliance",
                              "analyste risques financiers CDI"],
        "🌍 Luxembourg AML": ["AML officer Luxembourg CDI", "compliance Luxembourg",
                               "financial crime Luxembourg", "KYC analyst Luxembourg"],
        "🏛️ Asset Management": ["compliance asset management", "analyste conformité gestion actifs",
                                  "risk compliance fund management"],
        "🔒 Sanctions / Embargos": ["sanctions analyst CDI", "analyste sanctions OFAC",
                                     "screening sanctions compliance", "embargoes analyst"],
    }

    selected_cats = []
    cols_kw = st.columns(4)
    cat_names = list(kw_categories.keys())
    for i, cat in enumerate(cat_names):
        with cols_kw[i % 4]:
            if st.checkbox(cat, value=(i < 4), key=f"cat_{i}"):
                selected_cats.append(cat)

    # Build final keywords list
    all_keywords = []
    for cat in selected_cats:
        all_keywords.extend(kw_categories[cat])
    if custom_kw.strip():
        all_keywords.extend([kw.strip() for kw in custom_kw.split(",")])
    all_keywords = list(dict.fromkeys(all_keywords))  # deduplicate

    st.markdown(f"""
    <div style='background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.15);
                border-radius:8px;padding:10px 14px;margin:8px 0 16px;font-size:0.8rem;'>
      <b style='color:#C9A84C;'>🏷️ Mots-clés actifs ({len(all_keywords)}) :</b>
      {' · '.join(f'<span style="opacity:0.8;">{k}</span>' for k in all_keywords[:20])}
      {'...' if len(all_keywords) > 20 else ''}
    </div>
    """, unsafe_allow_html=True)

    st.divider()

    # ── REAL-TIME SEARCH ─────────────────────────────────────────────────────
    if launch_search:
        if not api_key:
            st.warning("⚠️ **Clé API Anthropic requise** pour le scraping en temps réel. Renseignez-la dans la barre latérale gauche.")
            st.info("💡 Sans clé API, utilisez les liens directs ci-dessous.")
        else:
            with st.spinner("🔍 Scraping en cours — LinkedIn, Indeed, France Travail, APEC, Jobs.lu... (30-60 secondes)"):
                jobs = _search_real_jobs(api_key, all_keywords, job_location)
                st.session_state["scraped_jobs"] = jobs
                st.session_state["scraped_location"] = job_location
                st.session_state["scraped_kw"] = all_keywords[:10]

    # ── DISPLAY SCRAPED RESULTS ───────────────────────────────────────────────
    if st.session_state.get("scraped_jobs"):
        jobs = st.session_state["scraped_jobs"]
        loc_used = st.session_state.get("scraped_location", "")
        kw_used  = st.session_state.get("scraped_kw", [])

        st.markdown(f"""
        <div style='background:rgba(46,204,113,0.1);border:1px solid rgba(46,204,113,0.3);
                    border-radius:10px;padding:12px 16px;margin-bottom:16px;'>
          <b style='color:#2ECC71;'>✅ {len(jobs)} offres trouvées</b>
          <span style='font-size:0.82rem;opacity:0.7;'> — {loc_used} — mots-clés: {", ".join(kw_used[:5])}</span>
          <span style='float:right;font-size:0.75rem;opacity:0.5;'>Scraping temps réel</span>
        </div>
        """, unsafe_allow_html=True)

        # Filter bar
        filter_txt = st.text_input("🔍 Filtrer les résultats", placeholder="ex: Luxembourg, CAMS, Fintech...",
                                    key="filter_results", label_visibility="collapsed")

        c1, c2 = st.columns(2)
        display_jobs = [j for j in jobs if not filter_txt or
                       filter_txt.lower() in (j.get("title","") + j.get("company","") +
                                               j.get("location","") + j.get("description","")).lower()]

        for i, job in enumerate(display_jobs):
            title   = job.get("title","Poste non précisé")
            company = job.get("company","")
            loc_j   = job.get("location","")
            url_j   = job.get("url","")
            date_j  = job.get("date","")
            desc    = job.get("description","")

            with (c1 if i%2==0 else c2):
                btn_html = (f'<a href="{url_j}" target="_blank" style="display:inline-block;'
                            f'background:linear-gradient(135deg,#C9A84C,#E8C97A);color:#0D1B2A;'
                            f'padding:5px 14px;border-radius:5px;font-size:0.76rem;font-weight:700;'
                            f'text-decoration:none;margin-top:6px;">Postuler →</a>'
                            if url_j else "")
                st.markdown(f"""
                <div style="background:var(--navy-mid);border:1px solid rgba(255,255,255,0.07);
                            border-left:3px solid #C9A84C;border-radius:0 8px 8px 0;
                            padding:12px 14px;margin-bottom:10px;">
                  <div style="font-size:0.88rem;font-weight:600;color:#F5F0E8;margin-bottom:3px;">
                    {title}
                  </div>
                  <div style="font-size:0.78rem;color:#C9A84C;margin-bottom:2px;">
                    {company}{f' · {loc_j}' if loc_j else ''}{f' · {date_j}' if date_j else ''}
                  </div>
                  <div style="font-size:0.78rem;opacity:0.7;line-height:1.4;margin-bottom:4px;">
                    {desc[:120] + '...' if len(desc)>120 else desc}
                  </div>
                  {btn_html}
                </div>""", unsafe_allow_html=True)

        # Export results
        if display_jobs:
            import json
            st.download_button(
                "📥 Exporter les offres (.json)",
                data=json.dumps(display_jobs, ensure_ascii=False, indent=2),
                file_name="offres_emploi_compliance.json",
                mime="application/json",
                key="export_jobs"
            )

    st.divider()

    # ── STATIC QUICK ACCESS (always visible) ─────────────────────────────────
    st.markdown("### 🔗 Accès Direct aux Plateformes — Recherches pré-configurées CDI")
    st.markdown("<p style='opacity:0.7;font-size:0.83rem;'>Liens directs avec requêtes optimisées pour votre profil LCB-FT/AML.</p>", unsafe_allow_html=True)

    platforms = [
        {
            "name": "LinkedIn Jobs",
            "icon": "💼",
            "color": "#0A66C2",
            "links": [
                ("Analyste LCB-FT — CDI Île-de-France",
                 "https://www.linkedin.com/jobs/search/?keywords=analyste+LCB-FT+AML&location=%C3%8Ele-de-France&f_JT=F&sortBy=DD&f_TPR=r604800"),
                ("Compliance Officer AML — CDI France",
                 "https://www.linkedin.com/jobs/search/?keywords=compliance+officer+securite+financiere+AML&location=France&f_JT=F&sortBy=DD"),
                ("AML/KYC Analyst — Luxembourg CDI",
                 "https://www.linkedin.com/jobs/search/?keywords=AML+KYC+analyst+compliance&location=Luxembourg&f_JT=F&sortBy=DD"),
                ("Chargé Conformité Fintech — CDI",
                 "https://www.linkedin.com/jobs/search/?keywords=charge+conformite+fintech+LCB-FT&location=France&f_JT=F&sortBy=DD"),
                ("Sanctions Analyst / Screening — CDI",
                 "https://www.linkedin.com/jobs/search/?keywords=sanctions+analyst+screening+compliance&location=France&f_JT=F&sortBy=DD"),
            ]
        },
        {
            "name": "France Travail",
            "icon": "🇫🇷",
            "color": "#E63946",
            "links": [
                ("Analyste LCB-FT Sécurité Financière",
                 "https://candidat.francetravail.fr/offres/recherche?motsCles=analyste+LCB-FT+s%C3%A9curit%C3%A9+financi%C3%A8re&typeContrat=CDI&lieux=75&sort=1"),
                ("Compliance Officer AML",
                 "https://candidat.francetravail.fr/offres/recherche?motsCles=compliance+officer+AML+conformit%C3%A9&typeContrat=CDI&lieux=75&sort=1"),
                ("KYC / Due Diligence",
                 "https://candidat.francetravail.fr/offres/recherche?motsCles=KYC+due+diligence+conformit%C3%A9&typeContrat=CDI&sort=1"),
                ("Analyste Crédit Risques",
                 "https://candidat.francetravail.fr/offres/recherche?motsCles=analyste+cr%C3%A9dit+risques+financi%C3%A8re&typeContrat=CDI&lieux=75&sort=1"),
            ]
        },
        {
            "name": "Welcome to the Jungle",
            "icon": "🌿",
            "color": "#41B89C",
            "links": [
                ("Compliance AML/LCB-FT CDI",
                 "https://www.welcometothejungle.com/fr/jobs?query=compliance+AML+LCB-FT&refinementList%5Bcontract_type_names%5D%5B%5D=CDI&sortBy=publishedAt_desc"),
                ("Sécurité Financière CDI",
                 "https://www.welcometothejungle.com/fr/jobs?query=s%C3%A9curit%C3%A9+financi%C3%A8re&refinementList%5Bcontract_type_names%5D%5B%5D=CDI&sortBy=publishedAt_desc"),
                ("Chargé Conformité Fintech",
                 "https://www.welcometothejungle.com/fr/jobs?query=charg%C3%A9+conformit%C3%A9+fintech&refinementList%5Bcontract_type_names%5D%5B%5D=CDI&sortBy=publishedAt_desc"),
            ]
        },
        {
            "name": "APEC",
            "icon": "📋",
            "color": "#003F88",
            "links": [
                ("Conformité AML/LCB-FT Cadre",
                 "https://www.apec.fr/candidat/recherche-emploi.html/emploi?motsCles=conformit%C3%A9+AML+LCB-FT&typeContrat=CDI&sortsBase=DATE_PUBLICATION&sortsOrder=DECREASE"),
                ("Analyste Risques Financiers",
                 "https://www.apec.fr/candidat/recherche-emploi.html/emploi?motsCles=analyste+risques+financiers&typeContrat=CDI&sortsBase=DATE_PUBLICATION&sortsOrder=DECREASE"),
                ("Analyste Crédit Senior",
                 "https://www.apec.fr/candidat/recherche-emploi.html/emploi?motsCles=analyste+cr%C3%A9dit+senior&typeContrat=CDI&sortsBase=DATE_PUBLICATION&sortsOrder=DECREASE"),
            ]
        },
        {
            "name": "Indeed France",
            "icon": "🔍",
            "color": "#2164F3",
            "links": [
                ("Analyste LCB-FT Paris CDI",
                 "https://fr.indeed.com/jobs?q=analyste+LCB-FT+s%C3%A9curit%C3%A9+financi%C3%A8re&l=Paris&jt=fulltime&sort=date"),
                ("Compliance Officer AML Paris",
                 "https://fr.indeed.com/jobs?q=compliance+officer+AML&l=Paris&jt=fulltime&sort=date"),
                ("Chargé Conformité Fintech",
                 "https://fr.indeed.com/jobs?q=charg%C3%A9+conformit%C3%A9+fintech+LCB-FT&l=France&jt=fulltime&sort=date"),
                ("KYC Analyst Île-de-France",
                 "https://fr.indeed.com/jobs?q=KYC+analyst+conformit%C3%A9&l=%C3%8Ele-de-France&jt=fulltime&sort=date"),
                ("Analyste Crédit Risques Paris",
                 "https://fr.indeed.com/jobs?q=analyste+cr%C3%A9dit+risques&l=Paris&jt=fulltime&sort=date"),
            ]
        },
        {
            "name": "Luxembourg",
            "icon": "🇱🇺",
            "color": "#EF3340",
            "links": [
                ("Compliance/AML Officer — Jobs.lu",
                 "https://www.jobs.lu/en/job-search/?search=compliance+AML+officer&sort=date&contract=permanent"),
                ("KYC/AML Analyst Luxembourg",
                 "https://www.jobs.lu/en/job-search/?search=KYC+AML+analyst&sort=date"),
                ("Financial Crime — Luxembourg",
                 "https://www.jobs.lu/en/job-search/?search=financial+crime+analyst&sort=date"),
                ("AML Compliance — Moovijob",
                 "https://www.moovijob.com/emploi?keywords=AML+conformite&country=lu&contract=cdi"),
            ]
        },
    ]

    for platform in platforms:
        st.markdown(f"""
        <div style='margin-bottom:8px;'>
          <div style='font-size:0.88rem;font-weight:700;color:{platform["color"]};
                      border-bottom:1px solid rgba(255,255,255,0.08);
                      padding:5px 0 3px;margin-bottom:6px;'>
            {platform["icon"]} {platform["name"]}
          </div>
        </div>""", unsafe_allow_html=True)

        col1, col2 = st.columns(2)
        for i, (title_l, url_l) in enumerate(platform["links"]):
            with (col1 if i%2==0 else col2):
                st.markdown(f"""
                <div style="background:var(--navy-mid);border:1px solid rgba(255,255,255,0.06);
                            border-left:3px solid {platform['color']};border-radius:0 7px 7px 0;
                            padding:9px 13px;margin-bottom:8px;">
                  <div style="font-size:0.83rem;color:#F5F0E8;margin-bottom:5px;">{title_l}</div>
                  <a href="{url_l}" target="_blank"
                     style="background:{platform['color']};color:white;padding:4px 12px;
                            border-radius:4px;font-size:0.74rem;font-weight:600;
                            text-decoration:none;display:inline-block;">
                    Voir les offres →
                  </a>
                </div>""", unsafe_allow_html=True)

    # ── KEYWORDS REFERENCE ────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 🏷️ Tous les mots-clés de recherche — Sécurité Financière")
    all_kws = [
        "LCB-FT","AML/CFT","KYC","KYB","Due Diligence","PPE","SOE","Tracfin",
        "Déclaration de soupçon","GAFI","ACPR","AMLD5","AMLD6","Examen Renforcé",
        "Compliance","Conformité bancaire","Sécurité financière","Financial Crime",
        "Sanctions","Embargos","OFAC","Screening","Gel des avoirs","PEP",
        "AML Analyst","Compliance Officer","Chargé de conformité","Analyste LCB-FT",
        "Risk Manager","Analyste crédit","Credit Analyst","Analyste risques",
        "Fintech compliance","PSP compliance","Cartographie des risques",
        "Beneficial owner","UBO","CDI","Asset Management","CIB","Corporate Banking",
        "CAMS","AMF","Actimize","Siron","Jura","Bloomberg","Looker",
    ]
    st.markdown("""<div style='display:flex;flex-wrap:wrap;gap:5px;'>""" +
        "".join(f'<span style="background:rgba(201,168,76,0.10);border:1px solid rgba(201,168,76,0.2);'
                f'border-radius:4px;padding:2px 9px;font-size:0.76rem;color:#E8C97A;">{k}</span>'
                for k in all_kws) +
        "</div>", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 10 — ASSISTANT IA CARRIÈRE (18 prompts)
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "🤖  Assistant IA Carrière":
    st.markdown('<p class="section-header">🤖 Assistant IA — Stratégie Carrière Compliance</p>', unsafe_allow_html=True)
    st.markdown("""<p style='opacity:0.75;margin-bottom:20px;line-height:1.7;max-width:720px;'>
    18 outils IA pour optimiser votre recherche d'emploi en sécurité financière —
    de la stratégie de ciblage à la négociation de package.
    </p>""", unsafe_allow_html=True)

    prompts = [
        ("Phase 1 — Stratégie & Positionnement", [
            ("🏛️ Audit de Gouvernance",
             "Analyse les rapports annuels et sanctions ACPR/OFAC d'une banque cible et propose une optimisation de leur dispositif LCB-FT.",
             "Analyse le dispositif de conformité de [Banque cible] : identifie leurs faiblesses en matière de screening ou classification des risques selon leurs dernières sanctions ACPR/OFAC. Propose une note d'optimisation d'une page avec recommandations concrètes."),
            ("📨 Approche directe MLRO",
             "Rédige un message LinkedIn pour le MLRO d'une banque, partageant une analyse de typologies de fraude crypto.",
             f"Rédige un message LinkedIn percutant pour le MLRO de [Entreprise]. Inclus une analyse concise d'une nouvelle typologie de fraude liée aux crypto-actifs ou au Trade-Based Money Laundering que j'ai identifiée. Mon profil : {st.session_state.get('edit_summary_fr','Analyste LCB-FT, expérience Delubac et HiPay, MBA ESLSCA, CAMS.')}."),
            ("📝 Optimisation CV par résultats",
             "Réécris les expériences en supprimant les tâches passives et en ajoutant des résultats chiffrés.",
             f"Agis comme recruteur expert en conformité bancaire. Réécris ces expériences pour un poste LCB-FT en supprimant les descriptions de tâches passives et en ajoutant des résultats chiffrés estimés :\n\nDelubac : {'; '.join(st.session_state.get('edit_exp',[{}])[0].get('bullets_fr',[])) if st.session_state.get('edit_exp') else ''}\n\nHiPay : {'; '.join(st.session_state.get('edit_exp',[{},{}])[1].get('bullets_fr',[])) if len(st.session_state.get('edit_exp',[])) > 1 else ''}"),
            ("💼 Simulation Entretien",
             "Génère les 5 cas de blanchiment les plus complexes pour un entretien chez [Banque] avec méthode STAR.",
             "Génère les 5 cas de blanchiment les plus complexes (schtroumpfage, sociétés écrans, TBML, financement terrorisme) pour un poste d'Analyste LCB-FT chez [Banque]. Détaille les étapes d'investigation avec la méthode STAR (Situation, Tâche, Action, Résultat)."),
        ]),
        ("Phase 2 — Analyse & Influence", [
            ("📊 Analyse Comparative RegTech",
             "Compare l'approche compliance de deux banques et comment l'IA peut automatiser la remédiation KYC.",
             "Compare l'approche de la Sécurité Financière de [Banque A] vs [Banque B]. Produis un document de 5 points expliquant comment l'intégration de l'IA (RegTech) pourrait automatiser leur remédiation KYC tout en restant conforme à la 6ème directive européenne (AMLD6)."),
            ("🎯 Pitch Compliance Manager",
             "Traduis une procédure de gel des avoirs en pitch 2 min pour convaincre un manager non-spécialiste.",
             "Traduis la procédure de gel des avoirs ou de détection de fraude en un pitch de 2 minutes pour convaincre un manager opérationnel non-spécialiste de l'importance de ce blocage sans ralentir le business. Utilise des exemples concrets d'amendes ACPR."),
            ("🧠 Matrice Questions Comportementales",
             "Crée une matrice des 10 questions comportementales en compliance avec mes expériences mappées.",
             f"Crée une matrice des 10 questions comportementales typiques en conformité (ex: 'Comment gérez-vous la pression métier pour valider un client douteux ?'). Mappe mes expériences :\n- Delubac : alertes N2, ER, DS/SAR\n- HiPay : PPE/SOE, pays sensibles, Tracfin\n- CAMS certifié\n- Hackathon Fintech 2023\nen mettant l'accent sur mon intégrité et ma gestion des priorités."),
            ("💰 Négociation Package",
             "Email de négociation +15% fixe justifié par la certification CAMS et la réduction des risques d'amendes.",
             "Rédige un email de négociation pour un poste d'analyste compliance demandant +15% sur la part fixe, justifié par ma certification CAMS (Sanctions & Embargos), mon AMF et ma capacité à réduire les risques d'amendes réglementaires pour la banque (exemples chiffrés d'amendes ACPR récentes)."),
        ]),
        ("Phase 3 — Skill-Up & Réseautage", [
            ("📚 Programme Auto-Formation 7j",
             "Programme de 7 jours pour maîtriser Actimize/Siron ou la réglementation MiCA.",
             "Crée un programme d'auto-formation de 7 jours pour maîtriser [Actimize / Siron / MiCA / DORA]. Inclus un exercice pratique d'analyse de transactions et les arguments à utiliser en entretien pour prouver ma montée en compétence. Adapte-le à mon profil LCB-FT."),
            ("🤝 Réactivation Réseau",
             "Message pour un ancien collègue perdu de vue via une actualité récente en compliance.",
             "Rédige un message pour un ancien collègue en conformité perdu de vue. Utilise une actualité récente (circulaire Banque de France, nouvelle directive GAFI, sanction ACPR récente) pour relancer l'échange professionnel avant d'évoquer le marché de l'emploi."),
            ("📋 Étude de Cas KYC/Sanctions",
             "Optimise mon raisonnement sur un cas KYC, identifie les UBO cachés et rédige une DS percutante.",
             "Examine cette étude de cas KYC/Sanctions : [décrire le cas]. Optimise mon raisonnement, identifie les bénéficiaires effectifs cachés (UBO) que j'aurais pu rater, détecte les signaux PPE/SOE, et rédige une conclusion de déclaration de soupçon (DS) percutante pour TRACFIN."),
            ("🏦 Adéquation Culture Éthique",
             "Analyse le code de conduite d'une banque et réécris mon profil pour résonner avec leurs valeurs.",
             "Analyse le code de conduite et les engagements RSE de [Banque cible]. Réécris mon 'About Me' LinkedIn pour qu'il résonne avec leurs valeurs de probité et de lutte contre la criminalité financière. Mon profil actuel : Analyste LCB-FT, CAMS, AMF, MBA ESLSCA."),
        ]),
        ("Phase 4 — Finalisation & Intégration", [
            ("🔄 Pivot après Rejet",
             "Email suite à une réponse négative pour comprendre les lacunes et rester en contact.",
             "Rédige un email suite à une réponse négative d'un recruteur en conformité. Demande diplomatiquement si mon analyse du risque ou ma connaissance réglementaire était insuffisante. Propose de rester en contact pour de futures missions de remédiation ou de renfort ponctuel."),
            ("📈 Optimisation LinkedIn",
             "Suggère 2 thématiques de rapports à publier pour démontrer mon expertise compliance.",
             f"Analyse mon profil LinkedIn (Analyste LCB-FT, expérience Delubac + HiPay, MBA Finance de Marché, CAMS, mémoire ESG). Suggère 2 thématiques de rapports/articles à publier (ex: analyse risque pays corridor Afrique/Europe, impact AMLD6 sur les PSP) pour démontrer mon expertise en sécurité financière."),
            ("🚀 Pitch Néo-Banque",
             "Email au Responsable Conformité d'une néo-banque pour diviser par deux le temps de revue des alertes.",
             "Rédige un email percutant au Responsable Conformité d'une néo-banque, expliquant comment ma double compétence en analyse de données (Python, Excel avancé, Looker) et LCB-FT peut diviser par deux le temps de revue des alertes grâce à l'automatisation des contrôles récurrents."),
            ("📅 Plan 30-60-90 Jours",
             "Plan pour les 90 premiers jours : maîtrise outils, revue stock, proposition d'amélioration monitoring.",
             "Génère un plan pour mes 90 premiers jours dans un poste Analyste LCB-FT senior : semaine 1 (maîtrise des outils internes Actimize/Siron, lecture procédures), mois 1 (revue complète du stock de dossiers en retard, KPIs), mois 3 (proposition d'amélioration des scénarios de monitoring, quick-wins chiffrés)."),
        ]),
    ]

    for phase_title, phase_prompts in prompts:
        st.markdown(f"""
        <div style='font-size:0.88rem;font-weight:700;color:#C9A84C;
                    border-bottom:1px solid rgba(201,168,76,0.2);
                    padding:8px 0 5px;margin:16px 0 10px;'>
          {phase_title}
        </div>""", unsafe_allow_html=True)

        for tool_title, tool_desc, tool_prompt in phase_prompts:
            with st.expander(f"**{tool_title}** — {tool_desc}"):
                st.markdown(f"""
                <div style='background:var(--navy-mid);border:1px solid rgba(201,168,76,0.15);
                            border-radius:8px;padding:12px;margin-bottom:10px;
                            font-size:0.85rem;color:var(--cream);opacity:0.9;line-height:1.6;
                            font-family:monospace;white-space:pre-wrap;'>
{tool_prompt}
                </div>""", unsafe_allow_html=True)
                col_copy, col_exec = st.columns(2)
                with col_copy:
                    st.code(tool_prompt[:200]+"…" if len(tool_prompt)>200 else tool_prompt,
                            language=None)
                with col_exec:
                    custom = st.text_area("Personnaliser le prompt", value=tool_prompt,
                                          height=80, key=f"prompt_{tool_title[:10]}",
                                          label_visibility="collapsed")
                    if st.button("🤖 Exécuter avec Claude", key=f"exec_{tool_title[:10]}"):
                        if not st.session_state.get("anthropic_api_key","").strip():
                            st.warning("⚠️ Clé API requise")
                        else:
                            with st.spinner("Claude travaille…"):
                                try:
                                    import requests as _req
                                    hdrs = {"Content-Type":"application/json",
                                            "x-api-key":st.session_state.get("anthropic_api_key",""),
                                            "anthropic-version":"2023-06-01"}
                                    body = {"model":"claude-sonnet-4-20250514","max_tokens":1000,
                                            "system":"Expert en conformité bancaire LCB-FT, compliance financière et stratégie de carrière. Réponds en français.",
                                            "messages":[{"role":"user","content":custom}]}
                                    r = _req.post("https://api.anthropic.com/v1/messages",headers=hdrs,json=body,timeout=60)
                                    result = "".join(b.get("text","") for b in r.json().get("content",[]))
                                    st.markdown(f"""
                                    <div style='background:rgba(201,168,76,0.08);border:1px solid rgba(201,168,76,0.2);
                                                border-radius:8px;padding:14px;margin-top:8px;
                                                font-size:0.88rem;line-height:1.7;white-space:pre-wrap;'>
                                    {result}
                                    </div>""", unsafe_allow_html=True)
                                except Exception as e:
                                    st.error(f"Erreur: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 11 — OUTILS COMPLIANCE
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "🔧  Outils Compliance":
    st.markdown('<p class="section-header">🔧 Outils Compliance & Ressources Réglementaires</p>', unsafe_allow_html=True)

    t1, t2, t3 = st.tabs(["📚 Réglementations clés","🌍 Pays GAFI","✅ Checklist KYC/AML"])

    with t1:
        regs = [
            ("AMLD5 (2018/843/UE)", "5e directive anti-blanchiment", "https://eur-lex.europa.eu/legal-content/FR/TXT/?uri=CELEX%3A32018L0843", "Renforcement KYC monnaies virtuelles, registres UBO, PPE"),
            ("AMLD6 (2018/1673/UE)", "6e directive — infractions sous-jacentes", "https://eur-lex.europa.eu/legal-content/FR/TXT/?uri=CELEX%3A32018L1673", "22 infractions prédicats, responsabilité pénale personnes morales"),
            ("Règlement DORA (2022)", "Digital Operational Resilience Act", "https://www.eba.europa.eu/regulation-and-policy/dora", "Résilience numérique des entités financières à partir de 2025"),
            ("Règlement MiCA (2023)", "Markets in Crypto-Assets", "https://www.eba.europa.eu/regulation-and-policy/mica", "Cadre réglementaire crypto-actifs UE"),
            ("Lignes directrices ACPR LCB-FT", "Guidelines ACPR 2023", "https://acpr.banque-france.fr/sites/default/files/media/2023/", "Approche par les risques, Examens Renforcés, cartographie"),
            ("40 Recommandations GAFI 2023", "Standards internationaux AML/CFT", "https://www.fatf-gafi.org/en/topics/fatf-recommendations.html", "Normes mondiales incluant actifs virtuels et VASP"),
        ]
        for reg, desc, url, detail in regs:
            st.markdown(f"""
            <div class="platform-card">
              <div class="platform-icon">📋</div>
              <div style="flex:1;">
                <div class="platform-name"><a href="{url}" target="_blank" style="color:#C9A84C;">{reg}</a></div>
                <div class="platform-desc">{desc} · {detail}</div>
              </div>
            </div>""", unsafe_allow_html=True)

    with t2:
        st.markdown("#### Pays à Surveillance Renforcée GAFI (extrait 2024)")
        col1, col2 = st.columns(2)
        pays_rouge = ["Iran","Corée du Nord","Myanmar","Russie (sanctions)",
                      "Cuba","Venezuela","Biélorussie"]
        pays_orange = ["Burkina Faso","Mali","Niger","Sénégal","Éthiopie",
                       "Nigeria","Cameroun","Mozambique","Tanzania","Kenya",
                       "Soudan du Sud","Syrie","Yémen","Liban"]
        with col1:
            st.markdown("🔴 **Liste noire / Appel à l'action**")
            for p in pays_rouge:
                st.markdown(f"• {p}")
        with col2:
            st.markdown("🟠 **Surveillance renforcée**")
            for p in pays_orange:
                st.markdown(f"• {p}")
        st.caption("Source: FATF/GAFI — listes régulièrement mises à jour. Vérifier sur fatf-gafi.org")

    with t3:
        st.markdown("#### Checklist KYC/AML Standard")
        checks = {
            "Identification du client": ["Pièce d'identité valide vérifiée","Adresse de résidence confirmée","Nationalité(s) documentée(s)"],
            "Bénéficiaire effectif (UBO)": ["Identification des UBO >25%","Documentation de la structure d'actionnariat","Vérification registre RBE (Infogreffe)"],
            "PPE / SOE Screening": ["Screening PPE (actuel et historique)","Vérification SOE (entreprises d'État)","Membres famille & proches associates"],
            "Sanctions & Listes": ["Screening OFAC SDN","Screening UE / CSFP","Screening ONU","Screening national (DGST/DGSI si pertinent)"],
            "Cohérence économique": ["Origine des fonds documentée","Cohérence revenus/transactions","Analyse du bilan si personne morale","Justificatifs opération si >15 000€"],
            "Décision finale": ["Niveau de risque attribué (faible/moyen/élevé)","ER si risque élevé","DS Tracfin si soupçon","Dossier archivé 5 ans minimum"],
        }
        for section_c, items in checks.items():
            st.markdown(f"**{section_c}**")
            for item in items:
                if f"chk_{item}" not in st.session_state:
                    st.session_state[f"chk_{item}"] = False
                st.checkbox(item, key=f"chk_{item}")
            st.markdown("")


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 12 — ÉDITION DU PROFIL
# ═══════════════════════════════════════════════════════════════════════════════
elif page == "✏️   Édition du Profil":
    st.markdown('<p class="section-header">✏️ Édition & Mise à Jour du Profil</p>', unsafe_allow_html=True)
    st.markdown("""<p style='opacity:0.75;margin-bottom:16px;'>
    Modifications appliquées <b style='color:#C9A84C;'>en temps réel</b> sur toute la plateforme et dans les PDFs.
    Cliquez <b>💾 Sauvegarder</b> pour valider.
    </p>""", unsafe_allow_html=True)

    t1,t2,t3,t4,t5,t6,t7,t8 = st.tabs([
        "👤 Identité","💼 Expériences","🎓 Formation",
        "⚡ Compétences","🏅 Distinctions","📖 Mémoire",
        "🧠 Personnalité","📝 Résumés FR/EN"
    ])

    with t1:
        c1,c2 = st.columns(2)
        with c1:
            st.text_input("Nom complet",       key="edit_name")
            st.text_input("Localisation",      key="edit_location")
            st.text_input("Email",             key="edit_email")
        with c2:
            st.text_input("Titre professionnel", key="edit_title")
            st.text_input("Téléphone",           key="edit_phone")
            st.text_input("LinkedIn",            key="edit_linkedin")
        if st.button("💾 Sauvegarder l'identité", use_container_width=True, key="save_id"):
            st.success("✅ Identité mise à jour sur toute la plateforme !")

    with t2:
        exps = st.session_state.edit_exp
        for i, exp in enumerate(exps):
            label_exp = f"**{exp['role']}** — {exp['org']} ({exp['period']})"
            with st.expander(label_exp, expanded=(i < 2)):
                c1,c2 = st.columns([2,1])
                with c1:
                    nr = st.text_input("Poste FR", value=exp.get("role",""),    key=f"er_{i}")
                    ne = st.text_input("Poste EN", value=exp.get("role_en",""), key=f"ere_{i}")
                    no = st.text_input("Organisation", value=exp.get("org",""),  key=f"eo_{i}")
                with c2:
                    np = st.text_input("Période FR", value=exp.get("period",""),    key=f"ep_{i}")
                    npe= st.text_input("Période EN", value=exp.get("period_en",""), key=f"epe_{i}")
                    nt = st.selectbox("Type", ["Banque","Fintech","Compétition","ONG","Association","Bénévolat"],
                                      key=f"et_{i}",
                                      index=["Banque","Fintech","Compétition","ONG","Association","Bénévolat"].index(
                                          exp.get("org_type","Banque")) if exp.get("org_type","Banque") in
                                          ["Banque","Fintech","Compétition","ONG","Association","Bénévolat"] else 0)
                st.markdown("**Missions FR** *(une par ligne)*")
                nbf = st.text_area("", value="\n".join(exp.get("bullets_fr",[])), height=100,
                                    key=f"ebf_{i}", label_visibility="collapsed")
                st.markdown("**Missions EN**")
                nbe = st.text_area("", value="\n".join(exp.get("bullets_en",[])), height=100,
                                    key=f"ebe_{i}", label_visibility="collapsed")
                col_s, col_d = st.columns([3,1])
                with col_s:
                    if st.button("💾 Sauvegarder", key=f"save_exp_{i}"):
                        st.session_state.edit_exp[i] = {
                            "role":nr,"role_en":ne,"org":no,"period":np,"period_en":npe,
                            "org_type":nt,
                            "bullets_fr":[b.strip() for b in nbf.split("\n") if b.strip()],
                            "bullets_en":[b.strip() for b in nbe.split("\n") if b.strip()],
                            "url":exp.get("url",""),
                        }
                        st.success(f"✅ « {nr} » sauvegardé !"); st.rerun()
                with col_d:
                    if st.button("🗑", key=f"del_exp_{i}", type="secondary"):
                        st.session_state.edit_exp.pop(i); st.rerun()

        st.markdown("---")
        with st.expander("➕ Nouvelle expérience"):
            c1,c2 = st.columns([2,1])
            with c1:
                new_r  = st.text_input("Poste FR",  key="new_er", placeholder="Analyste LCB-FT")
                new_re = st.text_input("Poste EN",  key="new_ere", placeholder="AML Analyst")
                new_o  = st.text_input("Org",       key="new_eo")
            with c2:
                new_p  = st.text_input("Période FR",key="new_ep")
                new_pe = st.text_input("Période EN",key="new_epe")
                new_t  = st.selectbox("Type",["Banque","Fintech","Compétition","ONG","Association"],key="new_et")
            new_bf = st.text_area("Missions FR",key="new_ebf",height=70)
            new_be = st.text_area("Missions EN",key="new_ebe",height=70)
            if st.button("➕ Ajouter",key="add_exp_btn"):
                if new_r and new_o:
                    st.session_state.edit_exp.insert(0,{
                        "role":new_r,"role_en":new_re,"org":new_o,"period":new_p,
                        "period_en":new_pe,"org_type":new_t,
                        "bullets_fr":[b.strip() for b in new_bf.split("\n") if b.strip()],
                        "bullets_en":[b.strip() for b in new_be.split("\n") if b.strip()],
                        "url":"",
                    })
                    st.success(f"✅ Ajouté !");  st.rerun()

    with t3:
        for i, e in enumerate(st.session_state.edit_edu):
            with st.expander(f"**{e['deg']}** — {e['school']}", expanded=(i==0)):
                c1,c2 = st.columns([2,1])
                with c1:
                    nd  = st.text_input("Diplôme FR",value=e.get("deg",""),  key=f"ed_{i}")
                    nde = st.text_input("Diplôme EN",value=e.get("deg_en",""),key=f"ede_{i}")
                    ns  = st.text_input("École",     value=e.get("school",""),key=f"es_{i}")
                with c2:
                    ny  = st.text_input("Années FR", value=e.get("yr",""),   key=f"ey_{i}")
                    nye = st.text_input("Années EN", value=e.get("yr_en",""),key=f"eye_{i}")
                    nm  = st.text_input("Mention",   value=e.get("mention",""),key=f"em_{i}")
                ndet  = st.text_area("Cours FR",value=e.get("det",""),   key=f"edet_{i}",height=70)
                ndete = st.text_area("Cours EN",value=e.get("det_en",""),key=f"edete_{i}",height=70)
                if st.button("💾 Sauvegarder",key=f"save_edu_{i}"):
                    st.session_state.edit_edu[i] = {
                        "deg":nd,"deg_en":nde,"school":ns,"yr":ny,"yr_en":nye,
                        "det":ndet,"det_en":ndete,"mention":nm,"url":e.get("url","")
                    }
                    st.success(f"✅ Sauvegardé !"); st.rerun()

    with t4:
        st.markdown("#### Compétences (glisser le curseur, cliquer 💾)")
        upd = {}
        for sk, lv in list(st.session_state.edit_skills.items()):
            c1,c2,c3 = st.columns([3,1,0.4])
            with c1: nsn = st.text_input("",value=sk,key=f"skn_{sk}",label_visibility="collapsed")
            with c2: nsl = st.number_input("",0,100,lv,key=f"skl_{sk}",label_visibility="collapsed")
            with c3:
                if st.button("🗑",key=f"skd_{sk}"):
                    d_sk = dict(st.session_state.edit_skills); d_sk.pop(sk,None)
                    st.session_state.edit_skills=d_sk; st.rerun()
            upd[nsn]=nsl
        c1,c2 = st.columns(2)
        with c1: new_sn=st.text_input("Nouvelle compétence",key="new_skn")
        with c2: new_sl=st.number_input("Niveau",0,100,80,key="new_skl")
        st.markdown("#### Certifications")
        certs_txt = "\n".join(st.session_state.get("edit_certifs",[]))
        new_certs = st.text_area("(une par ligne)",value=certs_txt,height=70,key="certs_ta")
        st.text_input("Outils informatiques",key="edit_tech")
        st.text_input("Centres d'intérêt",key="edit_interests")
        if st.button("💾 Tout sauvegarder",use_container_width=True,key="save_skills"):
            if new_sn.strip(): upd[new_sn.strip()]=new_sl
            st.session_state.edit_skills=upd
            st.session_state.edit_certifs=[l.strip() for l in new_certs.split("\n") if l.strip()]
            st.success("✅ Compétences et certifications sauvegardées !"); st.rerun()

    with t5:
        dists=list(st.session_state.edit_distinctions)
        for i,(ic,ti,de) in enumerate(dists):
            c1,c2,c3,c4=st.columns([0.4,1.8,3.5,0.4])
            with c1: ni=st.text_input("",value=ic,key=f"di_i_{i}",label_visibility="collapsed")
            with c2: nt=st.text_input("",value=ti,key=f"di_t_{i}",label_visibility="collapsed")
            with c3: nd=st.text_input("",value=de,key=f"di_d_{i}",label_visibility="collapsed")
            with c4:
                if st.button("🗑",key=f"di_del_{i}"):
                    st.session_state.edit_distinctions.pop(i); st.rerun()
        if st.button("💾 Sauvegarder distinctions",use_container_width=True,key="save_dists"):
            st.session_state.edit_distinctions=[(st.session_state.get(f"di_i_{i}",dists[i][0]),
                st.session_state.get(f"di_t_{i}",dists[i][1]),
                st.session_state.get(f"di_d_{i}",dists[i][2])) for i in range(len(dists))]
            st.success("✅ Sauvegardé !"); st.rerun()
        st.markdown("---")
        with st.expander("➕ Nouvelle distinction"):
            c1,c2,c3=st.columns([0.4,2,4])
            with c1: ndi=st.text_input("",key="ndi",value="🏅")
            with c2: ndt=st.text_input("Titre",key="ndt")
            with c3: ndd=st.text_input("Description",key="ndd")
            if st.button("➕",key="add_dist"):
                if ndt.strip():
                    st.session_state.edit_distinctions.append((ndi,ndt,ndd)); st.rerun()

    with t6:
        st.text_input("Titre du mémoire",key="edit_memoir_title",
                       placeholder="Impacts of ESG Factors…")
        st.text_area("Description",key="edit_memoir_desc",height=80)
        if st.button("💾 Sauvegarder mémoire",use_container_width=True,key="save_memoir"):
            st.success("✅ Mémoire sauvegardé — intégré dans CV et LM !")

    with t7:
        st.text_area("Phrase de personnalité",key="edit_personality",height=70)
        if st.button("💾 Sauvegarder",use_container_width=True,key="save_personality"):
            st.success("✅ Personnalité sauvegardée !")

    with t8:
        st.markdown("#### Résumé Profil (Français)")
        st.text_area("",key="edit_summary_fr",height=100,label_visibility="collapsed")
        st.markdown("#### Profile Summary (English)")
        st.text_area("",key="edit_summary_en",height=100,label_visibility="collapsed")
        if st.button("💾 Sauvegarder les résumés",use_container_width=True,key="save_summaries"):
            st.success("✅ Résumés sauvegardés — utilisés dans le CV FR et EN !")
